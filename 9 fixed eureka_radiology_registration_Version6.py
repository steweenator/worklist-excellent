# radiology_registration.py
# Full script, version with EVT_ACSE_REJECTED fix.

import sys
# print(f"DEBUG: Script using Python: {sys.executable}") # Uncomment for debug
# print(f"DEBUG: Python sys.path: {sys.path}") # Uncomment for debug

# --- Global Import Check for pydicom and pynetdicom ---
PYNETDICOM_AVAILABLE = False
# Initialize all potentially imported names to None, so they exist in the global scope
pydicom_module = None
Dataset = None
FileMetaDataset = None
AE = None
debug_logger = None # pynetdicom's debug_logger
evt = None
AllStoragePresentationContexts = None
ALL_TRANSFER_SYNTAXES = None
ModalityWorklistInformationFind = None
Verification = None # This is the SOP Class UID from pynetdicom.sop_class

# print("DEBUG: Top-level - Attempting pydicom and pynetdicom imports...") # Uncomment for debug
try:
    # Attempt all necessary imports
    import pydicom as pydicom_module_local
    from pydicom.dataset import Dataset as PydicomDatasetLocal, FileMetaDataset as PydicomFileMetaDatasetLocal
    from pynetdicom import (
        AE as PynetdicomAELocal,
        debug_logger as pynetdicom_debug_logger_local,
        evt as pynetdicom_evt_local,
        AllStoragePresentationContexts as PynetdicomAllStoragePresentationContextsLocal,
        ALL_TRANSFER_SYNTAXES as PynetdicomALL_TRANSFER_SYNTAXESLocal
    )
    from pynetdicom.sop_class import (
        ModalityWorklistInformationFind as PynetdicomModalityWorklistInformationFindLocal,
        Verification as PynetdicomVerificationLocal # Correct name for pynetdicom v2.x
    )

    # If all imports above were successful, assign them to the global variables
    pydicom_module = pydicom_module_local
    Dataset = PydicomDatasetLocal
    FileMetaDataset = PydicomFileMetaDatasetLocal
    AE = PynetdicomAELocal
    debug_logger = pynetdicom_debug_logger_local
    evt = pynetdicom_evt_local
    AllStoragePresentationContexts = PynetdicomAllStoragePresentationContextsLocal
    ALL_TRANSFER_SYNTAXES = PynetdicomALL_TRANSFER_SYNTAXESLocal
    ModalityWorklistInformationFind = PynetdicomModalityWorklistInformationFindLocal
    Verification = PynetdicomVerificationLocal # Assign the imported UID

    PYNETDICOM_AVAILABLE = True # Set flag to True ONLY if all imports succeeded
    # print("DEBUG: Top-level - pydicom and pynetdicom components LOADED SUCCESSFULLY.") # Uncomment for debug

except Exception as import_exception:
    # print(f"DEBUG: Top-level - FAILED to load pydicom/pynetdicom. Error: {import_exception}") # Uncomment for debug
    # import traceback # Uncomment for more detailed error
    # print("DEBUG: Traceback for import failure:") # Uncomment for more detailed error
    # traceback.print_exc() # Uncomment for more detailed error
    PYNETDICOM_AVAILABLE = False # Ensure flag is False on any import error
    # The global variables will remain None as initialized if imports fail
# --- End of Global Import Check ---

import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from datetime import datetime, timedelta
import sqlite3
# Enhanced python-docx import with debugging
import sys
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print(f"SUCCESS: python-docx imported successfully")
    print(f"Python executable: {sys.executable}")
except ImportError as e:
    print("WARNING: python-docx not available. DOCX generation will be disabled.")
    print(f"Import error: {e}")
    print(f"Python executable: {sys.executable}")
    print(f"Python version: {sys.version}")
    print("Install with: pip install python-docx")
    DOCX_AVAILABLE = False
    Document = None
import configparser
import shutil
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import re
import logging # Standard logging module
from watchdog.observers import Observer
from watchdog.events import PatternMatchingEventHandler
import threading

# --- Configuration and Constants ---
CONFIG_DIR = os.path.join(os.path.expanduser("~"), ".PatientRegistrationApp")
CONFIG_FILE = os.path.join(CONFIG_DIR, "config.ini")
DB_FILE = os.path.join(CONFIG_DIR, "patient_data.db")
DEFAULT_DATA_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "PatientRegistrationData")
LOG_FILE = os.path.join(CONFIG_DIR, "app.log")

if not os.path.exists(CONFIG_DIR):
    os.makedirs(CONFIG_DIR, exist_ok=True)

# Setup logging
logging.basicConfig(
    filename=LOG_FILE,
    level=logging.INFO, # Set to logging.DEBUG for more verbose output
    format="%(asctime)s [%(levelname)s] (%(module)s:%(lineno)d) %(message)s"
)
logging.info(f"--- Application Starting ---")
logging.info(f"Python Version: {sys.version.split()[0]}")
logging.info(f"Initial PYNETDICOM_AVAILABLE state after imports: {PYNETDICOM_AVAILABLE}")
logging.info(f"DOCX_AVAILABLE state: {DOCX_AVAILABLE}")
if PYNETDICOM_AVAILABLE:
    if pydicom_module:
        logging.info(f"pydicom version (if loaded): {getattr(pydicom_module, '__version__', 'N/A')}")
    if AE: 
        try:
            import pynetdicom as temp_pynetdicom
            logging.info(f"pynetdicom version (if loaded): {getattr(temp_pynetdicom, '__version__', 'N/A')}")
            del temp_pynetdicom
        except:
            logging.info("pynetdicom version could not be determined directly.")
else:
    logging.warning("PYNETDICOM_AVAILABLE is False. MWL Server functionality will be disabled.")


MODALITIES = ["CT", "DX", "US", "MG", "MR", "Default"]
SCANNER_SUPPORTED_FILE_TYPES = [".pdf", ".jpg", ".jpeg", ".png", ".tiff", ".bmp"]
EXTERNAL_REPORT_FILE_TYPES = (".pdf", ".doc", ".docx")
MODALITY_PATTERNS = ["*.pdf", "*.jpg", "*.jpeg", "*.png", "*.tiff", ".bmp"] 
EXTERNAL_REPORT_PATTERNS = ["*.pdf", "*.doc", ".docx"] 

# --- SQLite Database Helper Functions ---
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    
    # Check if we need to recreate the table with new schema
    cursor.execute("PRAGMA table_info(patient_records)")
    columns = [column[1] for column in cursor.fetchall()]
    
    # If old schema detected, backup and recreate
    if 'patient_id' in columns:
        try:
            # Check if patient_id is still unique constraint
            cursor.execute("SELECT sql FROM sqlite_master WHERE name='patient_records'")
            schema = cursor.fetchone()
            if schema and 'UNIQUE' in schema[0] and 'patient_id' in schema[0]:
                logging.info("Updating database schema to allow multiple studies per patient...")
                
                # Backup existing data
                cursor.execute("CREATE TABLE patient_records_backup AS SELECT * FROM patient_records")
                cursor.execute("DROP TABLE patient_records")
                
                # Create new table without unique constraints on patient_id
                cursor.execute('''
    CREATE TABLE patient_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_name TEXT NOT NULL,
        patient_id TEXT NOT NULL,
        accession_number TEXT NOT NULL UNIQUE,
        dob_yyyymmdd TEXT NOT NULL,
        sex TEXT NOT NULL,
        study_date TEXT NOT NULL,
        study_time TEXT NOT NULL,
        study_description TEXT NOT NULL,
        referred_from TEXT,
        modality TEXT NOT NULL,
        requesting_physician TEXT,
        requested_procedure_id TEXT,
        scheduled_station_ae_title TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
''')
                
                # Restore data (excluding removed fields)
                cursor.execute('''
    INSERT INTO patient_records 
    (patient_name, patient_id, accession_number, dob_yyyymmdd, sex,
     study_date, study_time, study_description, referred_from, modality,
     requesting_physician, requested_procedure_id, scheduled_station_ae_title, created_at)
    SELECT patient_name, patient_id, accession_number, dob_yyyymmdd, sex,
           study_date, study_time, study_description, referred_from, modality,
           requesting_physician, 
           COALESCE(requested_procedure_id, accession_number) as requested_procedure_id,
           COALESCE(scheduled_station_ae_title, 'ANY_MODALITY') as scheduled_station_ae_title,
           created_at
    FROM patient_records_backup
''')
                
                cursor.execute("DROP TABLE patient_records_backup")
                logging.info("Database schema updated successfully!")
        except Exception as e:
            logging.error(f"Error updating database schema: {e}")
    else:
        # Create new table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS patient_records (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_name TEXT NOT NULL,
                patient_id TEXT NOT NULL,
                accession_number TEXT NOT NULL UNIQUE,
                dob_yyyymmdd TEXT NOT NULL,
                sex TEXT NOT NULL,
                study_date TEXT NOT NULL,
                study_time TEXT NOT NULL,
                study_description TEXT NOT NULL,
                referred_from TEXT,
                modality TEXT NOT NULL,
                requesting_physician TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
    
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_patient_id ON patient_records (patient_id)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_accession_number ON patient_records (accession_number)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_study_date ON patient_records (study_date)")
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_modality ON patient_records (modality)")
    conn.commit()
    conn.close()
    logging.info(f"Database initialized/checked at {DB_FILE}")

def db_execute(query, params=(), fetchone=False, fetchall=False, commit=False):
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    try:
        cursor.execute(query, params)
        if commit:
            conn.commit()
            logging.debug(f"DB Execute (Commit): {query[:60]}... Params: {params}")
            return cursor.lastrowid if "INSERT" in query.upper() else True
        if fetchone:
            return cursor.fetchone()
        if fetchall:
            return cursor.fetchall()
        return True 
    except sqlite3.Error as e:
        logging.error(f"Database error: {e} \nQuery: {query} \nParams: {params}")
        return None
    finally:
        conn.close()

def add_patient_record_db(data_dict):
    query = '''
        INSERT INTO patient_records
        (patient_name, patient_id, accession_number, dob_yyyymmdd, sex,
         study_date, study_time, study_description, referred_from, modality,
         requesting_physician, requested_procedure_id, scheduled_station_ae_title)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    '''
    params = (
        data_dict.get("Patient Name"), data_dict.get("Patient ID"), data_dict.get("Accession Number"),
        data_dict.get("Date of Birth"), data_dict.get("Sex"), data_dict.get("Study Date"),
        data_dict.get("Study Time"), data_dict.get("Study Description"),
        data_dict.get("Referred From"), data_dict.get("Modality"),
        data_dict.get("Requesting Physician"), data_dict.get("Requested Procedure ID"),
        data_dict.get("Scheduled Station AE Title")
    )
    return db_execute(query, params, commit=True)

def update_patient_record_db(record_id, data_dict):
    """Update an existing patient record in the database"""
    query = '''
        UPDATE patient_records SET
        patient_name=?, patient_id=?, accession_number=?, dob_yyyymmdd=?, sex=?,
        study_date=?, study_time=?, study_description=?, referred_from=?, modality=?,
        requesting_physician=?
        WHERE id=?
    '''
    params = (
        data_dict.get("Patient Name"), data_dict.get("Patient ID"), data_dict.get("Accession Number"),
        data_dict.get("Date of Birth"), data_dict.get("Sex"), data_dict.get("Study Date"),
        data_dict.get("Study Time"), data_dict.get("Study Description"),
        data_dict.get("Referred From"), data_dict.get("Modality"),
        data_dict.get("Requesting Physician"),
        record_id
    )
    return db_execute(query, params, commit=True)

def get_patient_by_id_db(patient_id_to_find):
    query = "SELECT * FROM patient_records WHERE patient_id = ? ORDER BY created_at DESC LIMIT 1"
    row = db_execute(query, (patient_id_to_find,), fetchone=True)
    return dict(row) if row else None

def get_patient_record_by_db_id(record_id):
    """Get a specific patient record by database ID"""
    query = "SELECT * FROM patient_records WHERE id = ?"
    row = db_execute(query, (record_id,), fetchone=True)
    return dict(row) if row else None

def check_duplicate_record_db(patient_name, patient_id, accession_number):
    now = datetime.now()
    threshold_dt = now - timedelta(hours=36) 
    threshold_timestamp = threshold_dt.strftime("%Y-%m-%d %H:%M:%S")

    query_exact = """
        SELECT study_description, created_at FROM patient_records
        WHERE patient_id = ? AND accession_number = ? AND created_at > ?
        ORDER BY created_at DESC LIMIT 1
    """
    exact_match = db_execute(query_exact, (patient_id, accession_number, threshold_timestamp), fetchone=True)
    if exact_match:
        created_at_str = exact_match["created_at"]
        try:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S.%f")
        except ValueError:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S")
        diff = now - study_dt
        hrs = int(diff.total_seconds() // 3600)
        mins = int((diff.total_seconds() % 3600) // 60)
        return True, hrs, mins, study_dt.strftime("%b %d, %Y %H:%M:%S"), exact_match["study_description"], "Exact Patient ID and Accession match"

    query_general = """
        SELECT study_description, created_at FROM patient_records
        WHERE (patient_id = ? OR patient_name = ?) AND created_at > ?
        ORDER BY created_at DESC LIMIT 1
    """
    general_match = db_execute(query_general, (patient_id, patient_name, threshold_timestamp), fetchone=True)
    if general_match:
        created_at_str = general_match["created_at"]
        try:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S.%f")
        except ValueError:
            study_dt = datetime.strptime(created_at_str, "%Y-%m-%d %H:%M:%S")
        diff = now - study_dt
        hrs = int(diff.total_seconds() // 3600)
        mins = int((diff.total_seconds() % 3600) // 60)
        return True, hrs, mins, study_dt.strftime("%b %d, %Y %H:%M:%S"), general_match["study_description"], "Patient Name or ID match"

    return False, None, None, None, None, None 

def get_all_patient_records_db(search_term=""):
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row 
    cursor = conn.cursor()
    base_query_fields = "id, patient_name, patient_id, accession_number, dob_yyyymmdd, sex, study_date, study_time, study_description, referred_from, modality, requesting_physician, requested_procedure_id, scheduled_station_ae_title, created_at"
    if search_term:
        query = f"SELECT {base_query_fields} FROM patient_records WHERE patient_name LIKE ? OR patient_id LIKE ? OR accession_number LIKE ? OR study_description LIKE ? ORDER BY created_at DESC"
        like_term = f"%{search_term}%"
        cursor.execute(query, (like_term, like_term, like_term, like_term))
    else:
        query = f"SELECT {base_query_fields} FROM patient_records ORDER BY created_at DESC"
        cursor.execute(query)
    
    columns = [desc[0] for desc in cursor.description] if cursor.description else []
    all_data = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return all_data, columns

def get_distinct_values_for_combobox_db(field_name):
    rows = db_execute(f"SELECT DISTINCT {field_name} FROM patient_records WHERE {field_name} IS NOT NULL AND {field_name} != '' ORDER BY {field_name}", fetchall=True)
    return [row[0] for row in rows] if rows else []

def extract_patient_data_from_filename(filename):
    """Extract patient data from PDF filename similar to DOCX pattern parsing"""
    # Try to match the DOCX pattern: PatientName_Modality_Description_Date_REPORT.pdf
    match = re.match(r"^(.*?)_([A-Z0-9]+)_([^_]+)_([A-Za-z0-9_]+)_REPORT\.pdf$", filename, re.IGNORECASE)
    if match:
        patient_name = match.group(1).replace("_", " ")
        modality = match.group(2).upper()
        description = match.group(3).replace("_", " ")
        date_str = match.group(4).replace("_", " ")
        
        # Try to parse the date
        try:
            parsed_date = datetime.strptime(date_str, "%b %d %Y")
            display_date = parsed_date.strftime("%Y-%m-%d")
        except ValueError:
            display_date = date_str
        
        return {
            "patient_name": patient_name,
            "modality": modality,
            "description": description,
            "date": display_date,
            "filename": filename
        }
    
    # Fallback: try to extract at least some info from filename
    base_name = os.path.splitext(filename)[0]
    parts = base_name.split('_')
    
    if len(parts) >= 2:
        return {
            "patient_name": parts[0].replace("_", " "),
            "modality": parts[1] if len(parts) > 1 else "Unknown",
            "description": parts[2] if len(parts) > 2 else "Report",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "filename": filename
        }
    
    # Ultimate fallback
    return {
        "patient_name": "Unknown Patient",
        "modality": "Unknown",
        "description": "Report",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "filename": filename
    }


# --- MWL SCP Server Component ---
class MWLServerThread(threading.Thread):
    def __init__(self, app_config):
        super().__init__(daemon=True) 
        self.app_config = app_config
        self.ae_instance = None 
        self.server_running = False
        self.logger = logging.getLogger("mwl_scp") 

    def handle_echo(self, pynetdicom_event): 
        self.logger.info(f"C-ECHO request received from {pynetdicom_event.assoc.requestor.ae_title}@{pynetdicom_event.assoc.requestor.address}:{pynetdicom_event.assoc.requestor.port}")
        return 0x0000 

    def handle_find(self, pynetdicom_event): 
        self.logger.info(f"C-FIND request received from {pynetdicom_event.assoc.requestor.ae_title}@{pynetdicom_event.assoc.requestor.address}:{pynetdicom_event.assoc.requestor.port}")
        
        if not PYNETDICOM_AVAILABLE or Dataset is None or pydicom_module is None:
            self.logger.error("pydicom.Dataset or pydicom module not available, cannot process C-FIND.")
            yield 0xC001, None 
            return

        req_identifier = pynetdicom_event.identifier
        self.logger.debug(f"C-FIND Request Identifier:\n{req_identifier}")

        sql_query = "SELECT * FROM patient_records WHERE 1=1" 
        params = []

        if 'PatientName' in req_identifier and req_identifier.PatientName:
            pn = str(req_identifier.PatientName).replace('*', '%').replace('?', '_')
            sql_query += " AND patient_name LIKE ?"
            params.append(pn)
        if 'PatientID' in req_identifier and req_identifier.PatientID:
            sql_query += " AND patient_id = ?"
            params.append(str(req_identifier.PatientID))
        if 'AccessionNumber' in req_identifier and req_identifier.AccessionNumber:
            sql_query += " AND accession_number = ?"
            params.append(str(req_identifier.AccessionNumber))

        if 'ModalitiesInStudy' in req_identifier and req_identifier.ModalitiesInStudy : 
            sql_query += " AND modality = ?"
            params.append(str(req_identifier.ModalitiesInStudy))
        elif hasattr(req_identifier, 'ScheduledProcedureStepSequence') and \
             req_identifier.ScheduledProcedureStepSequence and \
             'Modality' in req_identifier.ScheduledProcedureStepSequence[0] and \
             req_identifier.ScheduledProcedureStepSequence[0].Modality: 
            sql_query += " AND modality = ?"
            params.append(str(req_identifier.ScheduledProcedureStepSequence[0].Modality))
        
        sps_start_date = ""
        if hasattr(req_identifier, 'ScheduledProcedureStepSequence') and \
           req_identifier.ScheduledProcedureStepSequence and \
           'ScheduledProcedureStepStartDate' in req_identifier.ScheduledProcedureStepSequence[0]:
            sps_start_date = req_identifier.ScheduledProcedureStepSequence[0].ScheduledProcedureStepStartDate

        if sps_start_date: 
            if '-' in sps_start_date: 
                start_date_range, end_date_range = sps_start_date.split('-')
                sql_query += " AND study_date BETWEEN ? AND ?"
                params.extend([start_date_range.strip(), end_date_range.strip()])
            else: 
                sql_query += " AND study_date = ?"
                params.append(sps_start_date.strip())
        
        self.logger.debug(f"Executing SQL for C-FIND: {sql_query} with params: {params}")
        matching_records = db_execute(sql_query, tuple(params), fetchall=True)

        if matching_records is None: 
            self.logger.error("Database error during C-FIND query.")
            yield 0xA700, None 
            return

        self.logger.info(f"Found {len(matching_records)} records matching C-FIND criteria.")

        for record in matching_records:
            if not PYNETDICOM_AVAILABLE or Dataset is None or pydicom_module is None: break 
            
            ds = Dataset() 

            ds.PatientName = record["patient_name"] if record["patient_name"] else "UNKNOWN"
            ds.PatientID = record["patient_id"] if record["patient_id"] else "UNKNOWN"
            ds.PatientBirthDate = record["dob_yyyymmdd"] if record["dob_yyyymmdd"] else "" 
            ds.PatientSex = record["sex"] if record["sex"] else "" 

            ds.AccessionNumber = record["accession_number"] if record["accession_number"] else ""
            ds.ReferringPhysicianName = record["referred_from"] if record["referred_from"] else "" 
            ds.StudyInstanceUID = pydicom_module.uid.generate_uid() 
            
            ds.RequestingPhysician = record["requesting_physician"] if record["requesting_physician"] else ""
            ds.RequestedProcedureDescription = record["study_description"] if record["study_description"] else "UNKNOWN"
            ds.RequestedProcedureID = record["requested_procedure_id"] if record["requested_procedure_id"] else ""

            sps_item = Dataset()
            sps_item.ScheduledStationAETitle = record["scheduled_station_ae_title"] if record["scheduled_station_ae_title"] else (self.ae_instance.ae_title if self.ae_instance else "UNKNOWN_AE")
            sps_item.ScheduledProcedureStepStartDate = record["study_date"] if record["study_date"] else "" 
            sps_item.ScheduledProcedureStepStartTime = record["study_time"] if record["study_time"] else "" 
            sps_item.Modality = record["modality"] if record["modality"] else ""
            sps_item.ScheduledPerformingPhysicianName = "" 
            sps_item.ScheduledProcedureStepDescription = record["study_description"] if record["study_description"] else "UNKNOWN"
            sps_item.ScheduledProcedureStepID = record["accession_number"] 

            ds.ScheduledProcedureStepSequence = [sps_item]

            if hasattr(req_identifier, 'SpecificCharacterSet'):
                ds.SpecificCharacterSet = req_identifier.SpecificCharacterSet
            else:
                ds.SpecificCharacterSet = "ISO_IR 100" 

            self.logger.debug(f"Yielding C-FIND response for Accession: {ds.AccessionNumber}")
            yield 0xFF00, ds 

        self.logger.info("Finished processing C-FIND, yielding final success status.")
        yield 0x0000, None 

    def _handle_acse_recv(self, event): # New handler for EVT_ACSE_RECV
        # This event is triggered when an A-ASSOCIATE-RJ or A-ABORT PDU is received.
        # The event object structure might differ slightly between pynetdicom versions.
        pdu_type_attr = None
        reason_diag_attr = 'N/A'
        result_attr = 'N/A'
        source_attr = 'N/A'
        
        # Try to determine PDU type and extract relevant info
        # pynetdicom 1.x style often has pdu_type directly
        if hasattr(event, 'pdu_type'): 
            pdu_type_attr = event.pdu_type
            reason_diag_attr = event.reason_diag if hasattr(event, 'reason_diag') else 'N/A'
            result_attr = event.result if hasattr(event, 'result') else 'N/A'
            source_attr = event.source if hasattr(event, 'source') else 'N/A'
        # pynetdicom 2.x style often uses event.primitive for A-ASSOCIATE-RJ
        elif hasattr(event, 'primitive') and hasattr(event.primitive, 'result_str'): # Check for result_str for A_ASSOCIATE_RJ
            pdu_type_attr = 0x03 # A-ASSOCIATE-RJ PDU type
            reason_diag_attr = event.primitive.information.get('Diagnostic', 'N/A') if hasattr(event.primitive, 'information') and isinstance(event.primitive.information, dict) else 'N/A'
            result_attr = event.primitive.result_str
            source_attr = event.primitive.source_str if hasattr(event.primitive, 'source_str') else 'N/A'
        # pynetdicom 2.x style for A-ABORT
        elif hasattr(event, 'primitive') and hasattr(event.primitive, 'source_str'): 
             pdu_type_attr = 0x07 # A-ABORT PDU type
             reason_diag_attr = event.primitive.information.get('Diagnostic', 'N/A') if hasattr(event.primitive, 'information') and isinstance(event.primitive.information, dict) else 'N/A'
             # No 'result' field in A-ABORT primitive
             source_attr = event.primitive.source_str
        else: 
            self.logger.info(f"ACSE PDU Received (structure unknown, cannot parse details): {event}")
            return

        ae_title_info = event.assoc.requestor.ae_title if event.assoc and event.assoc.requestor else 'Unknown AE'

        if pdu_type_attr == 0x03: # A-ASSOCIATE-RJ PDU
            self.logger.warning(f"Received A-ASSOCIATE-RJ (Rejection) from {ae_title_info}. Result: {result_attr}, Source: {source_attr}, Reason/Diag: {reason_diag_attr}")
        elif pdu_type_attr == 0x07: # A-ABORT PDU
             self.logger.error(f"Received A-ABORT PDU from {ae_title_info}. Source: {source_attr}, Reason/Diag: {reason_diag_attr}")
        else: # Other ACSE PDUs if any (shouldn't be common for this event)
            self.logger.info(f"ACSE PDU Received (Type {pdu_type_attr if pdu_type_attr else 'Unknown'}): {event}")


    def run(self):
        if not PYNETDICOM_AVAILABLE or AE is None or ModalityWorklistInformationFind is None or Verification is None or ALL_TRANSFER_SYNTAXES is None or evt is None:
            self.logger.error("Cannot start MWL SCP server: Essential pynetdicom components not available globally.")
            self.server_running = False
            return

        ae_title_str = self.app_config.get("MWLServerConfig", "ae_title", fallback="PYMWLSCP")
        port = self.app_config.getint("MWLServerConfig", "port", fallback=11112)

        try:
            self.ae_instance = AE(ae_title=ae_title_str.encode('ascii')) 
        except Exception as e_ae_init:
            self.logger.error(f"Failed to initialize AE for MWL Server: {e_ae_init}")
            self.server_running = False
            return
        
        self.ae_instance.add_supported_context(ModalityWorklistInformationFind, ALL_TRANSFER_SYNTAXES)
        self.ae_instance.add_supported_context(Verification, ALL_TRANSFER_SYNTAXES) 

        handlers = [
            (evt.EVT_C_ECHO, self.handle_echo),
            (evt.EVT_C_FIND, self.handle_find),
            (evt.EVT_ACSE_RECV, self._handle_acse_recv), # Corrected event and handler
            (evt.EVT_ABORTED, lambda event: self.logger.error(f"Connection Aborted: {event}")),
            (evt.EVT_REQUESTED, lambda event: self.logger.info(f"Association Requested: {event.assoc}")),
            (evt.EVT_ACCEPTED, lambda event: self.logger.info(f"Association Accepted by {event.assoc.acceptor.ae_title if event.assoc and event.assoc.acceptor else 'Unknown AE'}")),
            (evt.EVT_REJECTED, lambda event: self.logger.warning(f"Association Rejected by {event.assoc.acceptor.ae_title if event.assoc and event.assoc.acceptor else 'Unknown AE'}")), 
            (evt.EVT_RELEASED, lambda event: self.logger.info(f"Association Released: {event.assoc}")),
        ]

        self.logger.info(f"Starting MWL SCP server on port {port} with AE Title {ae_title_str}...")
        self.server_running = True
        try:
            self.ae_instance.start_server(('', port), block=True, evt_handlers=handlers)
        except OSError as e: 
            self.logger.error(f"OSError starting MWL SCP server (possibly port {port} in use): {e}")
            if tk._default_root and tk._default_root.winfo_exists():
                 messagebox.showerror("MWL Server Error", f"Could not start MWL server on port {port}.\nIs it already in use?\n\nError: {e}", parent=tk._default_root)
            else:
                 print(f"CRITICAL MWL Server Error (GUI not ready): Could not start MWL server on port {port}. Is it already in use? Error: {e}")
        except Exception as e: 
            self.logger.exception(f"General exception in MWL SCP server: {e}")
        finally:
            self.server_running = False
            self.logger.info("MWL SCP server has stopped.")

    def stop_server(self):
        if self.ae_instance and self.server_running:
            self.logger.info("Attempting to shut down MWL SCP server...")
            try:
                self.ae_instance.shutdown()
            except Exception as e_shutdown:
                self.logger.error(f"Error during MWL server shutdown: {e_shutdown}")
            self.server_running = False
            self.logger.info("MWL SCP server shutdown initiated.")
        else:
            self.logger.info("MWL SCP server was not running or AE instance not initialized.")


# --- WatchHandler Class (for file system monitoring) ---
class WatchHandler(PatternMatchingEventHandler):
    def __init__(self, patterns, callback, app_instance):
        super().__init__(patterns=patterns, ignore_directories=True, case_sensitive=False)
        self.callback = callback
        self.app = app_instance 

    def on_created(self, event):
        path = event.src_path
        logging.info(f"Watchdog detected new file via on_created: {path}")
        try:
            if self.app and hasattr(self.app, 'root') and self.app.root.winfo_exists():
                 self.app.root.after(500, lambda p=path: self._process_file(p))
            else:
                logging.warning(f"Root window not available/destroyed for delayed processing of {path}. Processing immediately.")
                self._process_file(path)
        except Exception as e: 
            logging.exception(f"Error scheduling callback for {path}: {e}. Processing immediately.")
            self._process_file(path)

    def _process_file(self, path):
        if os.path.exists(path) and os.path.getsize(path) > 0:
            logging.info(f"Processing confirmed new file: {path}")
            try:
                self.callback(path) 
            except Exception as e_callback:
                logging.exception(f"Error processing {path} in callback: {e_callback}")
        elif not os.path.exists(path):
            logging.warning(f"File {path} no longer exists, likely temporary. Skipping.")
        elif os.path.getsize(path) == 0:
            logging.warning(f"File {path} is empty, possibly still being written. Skipping for now.")

# --- Helper Functions ---
def get_script_directory():
    if getattr(sys, 'frozen', False): 
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__)) 

def ensure_dir_exists(path_to_ensure): 
    if path_to_ensure and not os.path.exists(path_to_ensure):
        try:
            os.makedirs(path_to_ensure, exist_ok=True)
            logging.info(f"Ensured directory exists: {path_to_ensure}")
        except Exception as e:
            logging.error(f"Failed to create directory {path_to_ensure}: {e}")

def normalize_path_for_config_section(path_str):
    if not path_str: return ""
    name = str(path_str).replace("\\", "_").replace("/", "_").replace(":", "_colon_").replace(" ", "_space_")
    name = re.sub(r'[^a-zA-Z0-9_.-]', '', name)
    return name

def load_config():
    config = configparser.ConfigParser(interpolation=None, allow_no_value=True)
    if not os.path.exists(CONFIG_FILE):
        logging.info(f"Config file not found at {CONFIG_FILE}. Creating default.")
        create_default_config(config) 
    else:
        try:
            config.read(CONFIG_FILE)
            logging.info(f"Config file loaded from {CONFIG_FILE}")
        except configparser.Error as e_cfg_read: 
            logging.exception(f"Error reading config file {CONFIG_FILE}: {e_cfg_read}. Will try to create default.")
            create_default_config(config)
        except Exception as e_generic_read: 
            logging.exception(f"Generic error reading config file {CONFIG_FILE}: {e_generic_read}. Will try to create default.")
            create_default_config(config)

    sections_to_check = [
        "Paths", "Preferences", "Paths.Output.DOCX.Modalities",
        "Paths.WatchFolders.Modalities", "UI.Labels", "SMTP",
        "Paths.ExternalReportWatchFolders", "EmailRecipients", "EmailTemplates",
        "MWLServerConfig"
    ]
    for section in sections_to_check:
        if not config.has_section(section):
            config.add_section(section)
            logging.info(f"Added missing section to config: {section}")

    default_db_path = DB_FILE
    default_template_path = os.path.join(DEFAULT_DATA_DIR, "Templates", "RADTEMPLATE.docx")
    default_general_docx_output = os.path.join(DEFAULT_DATA_DIR, "Reports", "_General")

    if not config.has_option("Paths", "db_file"): config.set("Paths", "db_file", default_db_path)
    if not config.has_option("Paths", "docx_template"): config.set("Paths", "docx_template", default_template_path)
    if not config.has_option("Paths", "general_docx_output_folder"): config.set("Paths", "general_docx_output_folder", default_general_docx_output)
    
    ensure_dir_exists(CONFIG_DIR)
    ensure_dir_exists(os.path.dirname(config.get("Paths", "db_file", fallback=default_db_path)))
    ensure_dir_exists(os.path.dirname(config.get("Paths", "docx_template", fallback=default_template_path)))
    ensure_dir_exists(config.get("Paths", "general_docx_output_folder", fallback=default_general_docx_output))
    
    return config

def save_config(config):
    ensure_dir_exists(CONFIG_DIR) 
    try:
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
        logging.info(f"Configuration saved to {CONFIG_FILE}")
    except Exception as e:
        logging.exception(f"Error writing config file {CONFIG_FILE} during save_config")

def create_default_config(config):
    logging.info("Creating default configuration file.")
    config.clear() 

    config['Paths'] = {
        'db_file': DB_FILE,
        'docx_template': os.path.join(DEFAULT_DATA_DIR, "Templates", "RADTEMPLATE.docx"),
        'general_docx_output_folder': os.path.join(DEFAULT_DATA_DIR, "Reports", "_General"),
        'general_watch_folder': "", 
        'external_report_watch_folders_list': "" 
    }
    config['Preferences'] = {
        'last_referred_from': '',
        'default_accession_prefix': 'CRH',
        'default_scheduled_station_ae': 'ANY_MODALITY',
        'color_theme': 'Default',
        'ui_style': 'System Default',
        'enable_tooltips': 'True',
        'ui_size': 'Default'
    }
    config['MWLServerConfig'] = {
        "enabled": "False",
        "ae_title": "PYREGMWL",
        "port": "11112"
    }
    config['Paths.Output.DOCX.Modalities'] = {}
    config['Paths.WatchFolders.Modalities'] = {}
    for mod in MODALITIES:
        config.set("Paths.Output.DOCX.Modalities", mod, os.path.join(DEFAULT_DATA_DIR, "Reports", mod))
        config.set("Paths.WatchFolders.Modalities", mod, os.path.join(DEFAULT_DATA_DIR, "WatchFolders", mod))

    default_labels = {
        "main_window_title": "Patient Registration & MWL Server", "patient_id": "Patient ID (e.g. 123456 AB):",
        "patient_name": "Patient Name:", "accession": "Accession (CRH[MODALITY]...):",
        "dob": "Date of Birth (DD/MM/YYYY or DDMMYYYY):", "sex": "Sex (M/F):",
        "study_description": "Study Description:", "referred_from": "Referred From:",
        "requesting_physician": "Requesting Physician:", "requested_procedure_id": "Requested Procedure ID:",
        "scheduled_station_ae": "Scheduled Station AE:",
        "attach_files_button": "Attach File(s)",
        "attachments_label_prefix": "Attachments:", "submit_button": "Register Patient", "clear_button": "Clear Form",
        "email_button": "Email Previous Report", "settings_window_title": "Settings",
        "view_data_window_title": "Patient Data Viewer",
        "view_served_worklist_title": "Served Worklist Viewer",
        "smtp_settings_tab_title": "SMTP Email",
        "smtp_server": "SMTP Server:", "smtp_port": "SMTP Port:", "smtp_user": "SMTP Username (optional):",
        "smtp_password": "SMTP Password (optional):", "smtp_sender_email": "Sender Email Address:",
        "smtp_use_tls": "Use TLS/STARTTLS Encryption", "smtp_test_button": "Test SMTP Settings",
        "email_picker_title": "Select Report to Email", "email_picker_button": "Compose Email for Selected Report",
        "email_composer_title": "Compose Email", "email_composer_to": "To:", "email_composer_subject": "Subject:",
        "email_composer_body": "Body:", "email_composer_attachments_label": "Attachments (auto-included):",
        "email_composer_send_button": "Send Email", "ext_reports_tab_title": "External Reports",
        "ext_reports_watch_folder_label": "Watch Folders for External Reports (PDF, DOC, DOCX):",
        "email_recipients_tab_title": "Email Recipients",
        "favorite_recipients_label": "Favorite Email Recipients:",
        "add_favorite_button": "Add Favorite",
        "remove_favorite_button": "Remove Selected Favorite",
        "max_recent_recipients_label": "Max Recent Recipients to Store:",
        "email_templates_tab_title": "Email Templates",
        "manage_email_templates_label": "Manage Email Templates:",
        "template_name_label": "Template Name:",
        "template_subject_label": "Subject Template:",
        "template_body_label": "Body Template:",
        "add_template_button": "Add New Template",
        "edit_template_button": "Edit Selected Template",
        "delete_template_button": "Delete Selected Template",
        "apply_template_button": "Apply Template",
        "available_placeholders_label": "Placeholders: {Patient Name}, {Modality}, {Study Description}, {Date}, {Report Filename}, {Attachment Count}, {All Attachment Names}",
        "email_composer_template_label": "Email Template:",
        "recent_recipients_combobox_label": "Recent/Favorites",
        "email_composer_select_recipients_label": "Select Recipients:",
        "email_composer_add_selected_button": "Add Selected to 'To:' Field",
        "ext_report_autosend_group_label": "Automatic Emailing for Selected Folder:",
        "ext_report_autosend_enable_label": "Enable Auto-Emailing for this Folder (PDF, DOC, DOCX):",
        "ext_report_autosend_recipients_label": "Auto-Send Recipient(s) (semicolon-separated):",
        "ext_report_autosend_add_favorite_button": "Add Favorite to Recipients",
        "ext_report_autosend_template_label": "Auto-Send Email Template:",
        "ext_report_autosend_custom_subject_label": "Custom Subject (if template is 'Custom'):",
        "ext_report_autosend_custom_body_label": "Custom Body (if template is 'Custom'):",
        "ext_report_autosend_placeholders_label": "Placeholders: {Filename}, {FolderPath}, {DateTime}",
        "appearance_tab_title": "Appearance Settings",
        "ui_style_engine_label": "UI Style Engine:",
        "color_palette_label": "Color Palette:",
        "app_mwl_server_tab_title": "This App as MWL Server",
        "app_mwl_server_enable_label": "Enable This Application as an MWL Server:",
        "app_mwl_server_ae_title_label": "This Application's AE Title:",
        "app_mwl_server_port_label": "Listening Port for DICOM Connections:"
    }
    config['UI.Labels'] = default_labels
    config['SMTP'] = {'server': '', 'port': '587', 'user': '', 'password': '', 'sender_email': '', 'use_tls': 'True'}
    config['Paths.ExternalReportWatchFolders'] = {} 
    config['EmailRecipients'] = {'recent_list': '', 'favorite_list': '', 'max_recent': '10'}
    
    default_email_subject = "Report: {Patient Name} - {Modality} {Study Description} on {Date}"
    default_email_body = ("Dear Recipient,\n\nPlease find attached the report for patient: {Patient Name}.\n"
                          "Study: {Modality} - {Study Description}\nDate: {Date}\n\n"
                          "Report file: {Report Filename}\n"
                          "Number of attached files: {Attachment Count}\n"
                          "Attached files: {All Attachment Names}\n\n"
                          "Best regards,\n[Your Name/Clinic]")
    config['EmailTemplates'] = {
        'template_names': 'Default', 
        'Default_subject': default_email_subject, 
        'Default_body': default_email_body
    }

    ensure_dir_exists(CONFIG_DIR)
    ensure_dir_exists(DEFAULT_DATA_DIR) 
    ensure_dir_exists(os.path.dirname(config.get("Paths", "db_file")))
    ensure_dir_exists(os.path.dirname(config.get("Paths", "docx_template")))
    ensure_dir_exists(config.get("Paths", "general_docx_output_folder"))
    
    general_watch_folder = config.get("Paths", "general_watch_folder", fallback="")
    if general_watch_folder: ensure_dir_exists(general_watch_folder)
    
    ext_report_folders_str = config.get("Paths", "external_report_watch_folders_list", fallback="")
    for folder in ext_report_folders_str.split(';'): 
        if folder: ensure_dir_exists(os.path.normpath(os.path.expanduser(folder)))
        
    for mod in MODALITIES:
        ensure_dir_exists(config.get("Paths.Output.DOCX.Modalities", mod))
        mod_watch_folder = config.get("Paths.WatchFolders.Modalities", mod)
        if mod_watch_folder: ensure_dir_exists(mod_watch_folder)

    default_template_filename = "RADTEMPLATE.docx"
    source_template_path = os.path.join(get_script_directory(), default_template_filename)
    target_template_path = config.get("Paths", "docx_template")
    if os.path.exists(source_template_path) and not os.path.exists(target_template_path):
        try:
            ensure_dir_exists(os.path.dirname(target_template_path)) 
            shutil.copy2(source_template_path, target_template_path)
            logging.info(f"Copied default template '{default_template_filename}' to {target_template_path}")
        except Exception as e:
            logging.exception(f"Could not copy default template: {e}")
    elif not os.path.exists(source_template_path):
        logging.warning(f"Default template '{default_template_filename}' not found in script directory: {get_script_directory()}. User needs to provide one.")
    
    save_config(config) 
    init_db() 
    logging.info(f"Default configuration file created and saved at {CONFIG_FILE}. Default data folders in {DEFAULT_DATA_DIR}. Database at {DB_FILE}")

def get_modality_from_accession(accession_number):
    acc_upper = accession_number.upper()
    if acc_upper.startswith("CRHCT"): return "CT"
    if acc_upper.startswith("CRHDX"): return "DX"
    if acc_upper.startswith("CRHUS"): return "US"
    if acc_upper.startswith("CRHMG"): return "MG"
    if acc_upper.startswith("CRHMR"): return "MR"
    return None 

def format_date_friendly(date_str_yyyymmdd):
    if not date_str_yyyymmdd: return "N/A"
    try:
        return datetime.strptime(date_str_yyyymmdd, "%Y%m%d").strftime("%b %d, %Y")
    except ValueError:
        return date_str_yyyymmdd 

# --- ToolTip Class ---
class ToolTip:
    def __init__(self, widget, text, app_config, main_app_instance):
        self.widget = widget
        self.text = text
        self.app_config = app_config
        self.main_app = main_app_instance 
        self.tooltip = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event=None):
        if not self.app_config.getboolean("Preferences", "enable_tooltips", fallback=True):
            return
        if self.tooltip: 
            self.tooltip.destroy()
        
        tooltip_bg = "#FFFFE0" 
        if hasattr(self.main_app, 'current_palette') and self.main_app.current_palette:
            tooltip_bg = self.main_app.current_palette.get("tooltip_bg", "#FFFFE0")

        x_root, y_root = self.widget.winfo_rootx(), self.widget.winfo_rooty()
        y_final = y_root + self.widget.winfo_height() + 5

        self.tooltip = tk.Toplevel(self.widget)
        self.tooltip.wm_overrideredirect(True) 
        
        label = ttk.Label(self.tooltip, text=self.text, 
                          background=tooltip_bg, 
                          relief="solid", borderwidth=1, padding=5,
                          wraplength=350) 
        label.pack()
        
        self.tooltip.update_idletasks() 
        tooltip_width = self.tooltip.winfo_width()
        
        final_x_pos = x_root + (self.widget.winfo_width() - tooltip_width) // 2
        
        screen_width = self.widget.winfo_screenwidth()
        if final_x_pos + tooltip_width > screen_width:
            final_x_pos = screen_width - tooltip_width - 10 
        if final_x_pos < 5 : 
            final_x_pos = 5
            
        self.tooltip.wm_geometry(f"+{int(final_x_pos)}+{int(y_final)}")

    def hide_tooltip(self, event=None):
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None

# --- PatientRegistrationApp Class ---
class PatientRegistrationApp:
    def __init__(self, root_window):
        self.root = root_window
        init_db() 
        self.config = load_config() 
        self.manual_attachments = []
        self.current_palette = {} 
        self.style = ttk.Style(self.root)

        self.main_frame = None 
        self.status_bar = None 
        self.mwl_server_thread = None 

        self.apply_theme_and_styles() 
        
        self.root.title(self.get_ui_label("main_window_title", "Patient Registration & MWL Server"))
        self.apply_ui_size()
        self.create_menu()

        self._selected_template_for_edit = None
        self.ext_report_autosend_widgets = {} 
        self._selected_ext_report_folder_for_autosend_config = None
        self.num_fields = 0 
        self._observer = None 

        self.create_widgets() 
        self.create_status_bar()
        self.load_combobox_values_from_db()

        if hasattr(self, 'entry_referred_from'):
            self.entry_referred_from.set(self.config.get("Preferences", "last_referred_from", fallback=""))
        if hasattr(self, 'entry_accession'):
            self.entry_accession.insert(0, self.config.get("Preferences", "default_accession_prefix", fallback="CRH"))
        
        if hasattr(self, 'entry_patient_id'): 
            self.entry_patient_id.focus()

        self._start_realtime_watchers()
        self.update_email_button_state() 
        self.start_mwl_server_if_configured()
        logging.info("PatientRegistrationApp initialized successfully.")

    def start_mwl_server_if_configured(self):
        if not PYNETDICOM_AVAILABLE:
            self.update_status("MWL Server disabled: pynetdicom/pydicom components not found.", True, 0)
            return

        if self.config.getboolean("MWLServerConfig", "enabled", fallback=False):
            if self.mwl_server_thread and self.mwl_server_thread.is_alive():
                logging.info("MWL Server thread already running.")
                return 

            self.mwl_server_thread = MWLServerThread(self.config)
            self.mwl_server_thread.start()
            self.root.after(1000, self.check_mwl_server_status)
        else:
            logging.info("MWL Server is disabled in configuration.")
            self.update_status("MWL Server is disabled.", False, 5000)

    def check_mwl_server_status(self):
        if not PYNETDICOM_AVAILABLE: 
            self.update_status("MWL Server cannot run: pynetdicom/pydicom missing.", True, 0)
            return

        if self.mwl_server_thread and self.mwl_server_thread.is_alive() and self.mwl_server_thread.server_running:
            ae_title = self.config.get("MWLServerConfig", "ae_title", fallback="N/A")
            port = self.config.get("MWLServerConfig", "port", fallback="N/A")
            self.update_status(f"MWL Server running: {ae_title} on port {port}", False, 0)
            logging.info(f"MWL Server confirmed running: {ae_title} on port {port}")
        elif self.config.getboolean("MWLServerConfig", "enabled", fallback=False):
            self.update_status("MWL Server failed to start or stopped. Check logs.", True, 0)
            logging.error("MWL Server was enabled but is not running. Check logs for errors (e.g., port conflict or pynetdicom issue).")

    def stop_mwl_server(self):
        if self.mwl_server_thread and self.mwl_server_thread.is_alive():
            logging.info("Stopping MWL Server thread...")
            self.mwl_server_thread.stop_server() 
            self.mwl_server_thread.join(timeout=5) 
            if self.mwl_server_thread.is_alive():
                logging.warning("MWL Server thread did not stop gracefully after 5 seconds.")
            else:
                logging.info("MWL Server thread stopped.")
        self.mwl_server_thread = None 
        self.update_status("MWL Server stopped.", False, 5000)

    def _start_realtime_watchers(self):
        if self._observer is not None: 
            try:
                self._observer.stop()
                self._observer.join(timeout=1) 
            except Exception as e:
                logging.warning(f"Could not stop previous Watchdog observer: {e}")
        
        self._observer = Observer()
        watched_folders = set() 

        for mod in MODALITIES:
            folder = self.get_modality_specific_path("Paths.WatchFolders.Modalities", mod)
            if folder and os.path.isdir(folder) and folder not in watched_folders:
                try:
                    handler = WatchHandler(MODALITY_PATTERNS, lambda p, m=mod: self._on_new_modality_file(p, m), self)
                    self._observer.schedule(handler, folder, recursive=False)
                    watched_folders.add(folder)
                    logging.info(f"Scheduled watchdog for modality folder ({mod}): {folder}")
                except Exception as e:
                    logging.error(f"Failed to schedule watchdog for modality folder {folder} ({mod}): {e}")
            elif folder and not os.path.isdir(folder):
                logging.warning(f"Modality watch folder for {mod} is not a valid directory: {folder}")

        gen_folder_path = self.config.get("Paths", "general_watch_folder", fallback="")
        if gen_folder_path and os.path.isdir(gen_folder_path) and gen_folder_path not in watched_folders:
            try:
                handler = WatchHandler(MODALITY_PATTERNS, lambda p: self._on_new_modality_file(p, None), self) 
                self._observer.schedule(handler, gen_folder_path, recursive=False)
                watched_folders.add(gen_folder_path)
                logging.info(f"Scheduled watchdog for general watch folder: {gen_folder_path}")
            except Exception as e:
                logging.error(f"Failed to schedule watchdog for general watch folder {gen_folder_path}: {e}")
        elif gen_folder_path and not os.path.isdir(gen_folder_path):
            logging.warning(f"General watch folder is not a valid directory: {gen_folder_path}")

        ext_report_folders = self.get_all_external_report_watch_folders()
        for folder in ext_report_folders:
            if folder and os.path.isdir(folder) and folder not in watched_folders:
                try:
                    handler = WatchHandler(EXTERNAL_REPORT_PATTERNS, self._process_auto_send_for_external_report_file, self)
                    self._observer.schedule(handler, folder, recursive=False)
                    watched_folders.add(folder)
                    logging.info(f"Scheduled watchdog for external report folder: {folder}")
                except Exception as e:
                    logging.error(f"Failed to schedule watchdog for external report folder {folder}: {e}")
            elif folder and not os.path.isdir(folder):
                 logging.warning(f"External report watch folder is not a valid directory: {folder}")

        if not self._observer.emitters: 
            logging.warning("Watchdog observer has no paths to watch. No watch folders configured or valid.")
        else:
            try:
                self._observer.start()
                logging.info("Watchdog observer started successfully.")
            except Exception as e:
                logging.exception("Failed to start Watchdog observer.")

    def _on_new_modality_file(self, filepath, modality_code):
        if not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
            logging.warning(f"File {filepath} is empty or gone, not staging for attachments.")
            return

        fname = os.path.basename(filepath)
        logging.info(f"Processing new modality file: {fname} (modality={modality_code or 'General'})")
        
        # Stage for attachment (existing functionality)
        if filepath not in self.manual_attachments:
            self.manual_attachments.append(filepath)
            self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")
            logging.info(f"File {fname} added to manual_attachments. Total: {len(self.manual_attachments)}")
        
        # NEW: Also copy to the modality output folder
        if modality_code:
            modality_output_folder = self.get_modality_specific_path("Paths.Output.DOCX.Modalities", modality_code)
            if modality_output_folder and os.path.isdir(modality_output_folder):
                try:
                    dest_path = os.path.join(modality_output_folder, fname)
                    if not os.path.exists(dest_path):  # Don't overwrite existing files
                        shutil.copy2(filepath, dest_path)
                        self.update_status(f"Auto-copied {fname} to {modality_code} output folder", duration=4000)
                        logging.info(f"Auto-copied {fname} from watch folder to output folder: {dest_path}")
                    else:
                        logging.info(f"File {fname} already exists in output folder, skipping copy")
                except Exception as e:
                    logging.error(f"Error copying {fname} to output folder: {e}")
                    self.update_status(f"Error copying {fname}: {e}", True, duration=6000)
            else:
                logging.warning(f"No valid output folder configured for modality {modality_code}")
        else:
            # For general folder, copy to general output folder
            general_output = self.config.get("Paths", "general_docx_output_folder", fallback="")
            if general_output and os.path.isdir(general_output):
                try:
                    dest_path = os.path.join(general_output, fname)
                    if not os.path.exists(dest_path):  # Don't overwrite existing files
                        shutil.copy2(filepath, dest_path)
                        self.update_status(f"Auto-copied {fname} to general output folder", duration=4000)
                        logging.info(f"Auto-copied {fname} from general watch folder to general output: {dest_path}")
                    else:
                        logging.info(f"File {fname} already exists in general output folder, skipping copy")
                except Exception as e:
                    logging.error(f"Error copying {fname} to general output folder: {e}")
                    self.update_status(f"Error copying {fname}: {e}", True, duration=6000)
            
    def shutdown(self):
        logging.info("Application shutdown sequence initiated.")
        self.stop_mwl_server() 

        if self._observer and self._observer.is_alive():
            try:
                self._observer.stop()
                self._observer.join(timeout=2) 
                logging.info("Watchdog observer stopped.")
            except Exception as e:
                logging.exception("Exception during Watchdog observer shutdown.")
        else:
            logging.info("Watchdog observer was not running or already stopped.")
        
        logging.info("Application shutdown complete.")

    def get_ui_label(self, key, default_text=""):
        return self.config.get("UI.Labels", key, fallback=default_text)

    def apply_ui_size(self):
        size_setting = self.config.get("Preferences", "ui_size", fallback="Default")
        sizes = {
            "Very Compact": "700x580", 
            "Compact": "700x610", 
            "Default": "750x680", 
            "Large": "850x750", 
            "Extra Large": "950x800"
        }
        self.root.geometry(sizes.get(size_setting, sizes["Default"]))
        logging.debug(f"UI size set to: {sizes.get(size_setting, sizes['Default'])}")

    def apply_theme_and_styles(self):
        selected_theme_name = self.config.get("Preferences", "color_theme", fallback="Default")
        selected_ui_style = self.config.get("Preferences", "ui_style", fallback="System Default")
        logging.info(f"Applying UI Style: {selected_ui_style}, Color Theme: {selected_theme_name}")

        base_ttk_theme = "clam" 
        if selected_ui_style == "Clam (Modern)": base_ttk_theme = "clam"
        elif selected_ui_style == "Alt (Modern-ish)": base_ttk_theme = "alt"
        elif selected_ui_style == "Default (Classic)": base_ttk_theme = "default" 
        elif selected_ui_style == "Classic (Older)": base_ttk_theme = "classic"
        elif selected_ui_style == "System Default":
            available_themes = self.style.theme_names()
            if 'vista' in available_themes and sys.platform == "win32": base_ttk_theme = 'vista'
            elif 'aqua' in available_themes and sys.platform == "darwin": base_ttk_theme = 'aqua' 
            elif 'clam' in available_themes: base_ttk_theme = 'clam' 
            else: base_ttk_theme = self.style.theme_use() 

        try:
            self.style.theme_use(base_ttk_theme)
            logging.info(f"Applied base TTK theme: {base_ttk_theme}")
        except tk.TclError:
            logging.warning(f"TTK theme '{base_ttk_theme}' not found, trying 'clam' as fallback.")
            try:
                base_ttk_theme = "clam"
                self.style.theme_use(base_ttk_theme)
            except tk.TclError:
                current_theme = self.style.theme_use()
                logging.error(f"Critical: TTK theme 'clam' also unavailable. Using system default: {current_theme}")
        
        themes = {
            "Default": {"bg": "#F0F0F0", "fg": "black", "entry_bg": "white", "entry_fg": "black", "button_bg": "#E0E0E0", "button_fg": "black", "label_fg": "black", "frame_bg": "#F0F0F0", "header_fg": "#0078D7", "status_bg": "#F0F0F0", "button_active_bg": "#D0D0D0", "tooltip_bg": "#FFFFE0"},
            "Light Blue": {"bg": "#E6F3FF", "fg": "#003366", "entry_bg": "#FFFFFF", "entry_fg": "#003366", "button_bg": "#B3D9FF", "button_fg": "#003366", "label_fg": "#004C99", "frame_bg": "#E6F3FF", "header_fg": "#003366", "status_bg": "#E6F3FF", "button_active_bg": "#80C7FF", "tooltip_bg": "#F0F8FF"},
            "Dark": {"bg": "#2E2E2E", "fg": "#E0E0E0", "entry_bg": "#3C3C3C", "entry_fg": "#E0E0E0", "button_bg": "#505050", "button_fg": "#FFFFFF", "label_fg": "#C0C0C0", "frame_bg": "#2E2E2E", "header_fg": "#4A9EFF", "status_bg": "#2E2E2E", "button_active_bg": "#606060", "tooltip_bg": "#404040"},
            "High Contrast": {"bg": "white", "fg": "black", "entry_bg": "white", "entry_fg": "black", "button_bg": "black", "button_fg": "white", "label_fg": "black", "frame_bg": "white", "header_fg": "black", "status_bg": "white", "button_active_bg": "#333333", "tooltip_bg": "#FFFFCC"},
            "Mint Green": {"bg": "#E0F2F1", "fg": "#004D40", "entry_bg": "#FFFFFF", "entry_fg": "#004D40", "button_bg": "#A7FFEB", "button_fg": "#004D40", "label_fg": "#00695C", "frame_bg": "#E0F2F1", "header_fg": "#004D40", "status_bg": "#E0F2F1", "button_active_bg": "#7FFFD4", "tooltip_bg": "#F0FFF0"},
            "Lavender": {"bg": "#F3E5F5", "fg": "#4A148C", "entry_bg": "#FFFFFF", "entry_fg": "#4A148C", "button_bg": "#E1BEE7", "button_fg": "#4A148C", "label_fg": "#6A1B9A", "frame_bg": "#F3E5F5", "header_fg": "#4A148C", "status_bg": "#F3E5F5", "button_active_bg": "#DDA0DD", "tooltip_bg": "#F8F0FF"}
        }
        palette = themes.get(selected_theme_name, themes["Default"])
        self.current_palette = palette 

        self.root.configure(bg=palette["bg"])
        self.style.configure('.', background=palette["bg"], foreground=palette["fg"]) 
        
        self.style.configure('TLabel', font=('Helvetica', 11), padding=5, background=palette["bg"], foreground=palette["label_fg"])
        self.style.configure('TButton', font=('Helvetica', 11, 'bold'), padding=5, background=palette["button_bg"], foreground=palette["button_fg"])
        self.style.map('TButton', background=[('active', palette.get("button_active_bg", palette["button_bg"]))]) 
        
        self.style.configure('TEntry', font=('Helvetica', 11), padding=5) 
        self.style.map('TEntry', fieldbackground=[('!focus', palette["entry_bg"]), ('focus', palette["entry_bg"])], foreground=[('!focus', palette["entry_fg"]), ('focus', palette["entry_fg"])])

        self.style.configure('TCombobox', font=('Helvetica', 11), padding=5)
        self.style.map('TCombobox', fieldbackground=[('readonly', palette["entry_bg"]), ('!readonly', palette["entry_bg"])], foreground=[('readonly', palette["entry_fg"]), ('!readonly', palette["entry_fg"])])
        self.root.option_add('*TCombobox*Listbox.background', palette["entry_bg"])
        self.root.option_add('*TCombobox*Listbox.foreground', palette["entry_fg"])
        self.root.option_add('*TCombobox*Listbox.selectBackground', palette.get("header_fg", "#0078D7")) 
        self.root.option_add('*TCombobox*Listbox.selectForeground', palette.get("button_fg", "white"))

        self.style.configure('Header.TLabel', font=('Helvetica', 14, 'bold'), foreground=palette["header_fg"], background=palette["bg"])
        self.style.configure('Custom.TFrame', background=palette["frame_bg"]) 
        self.style.configure('Status.TLabel', background=palette.get("status_bg", palette["bg"]), foreground=palette["fg"])

        self.style.configure('Treeview.Heading', font=('Helvetica', 10, 'bold'), background=palette.get("button_bg"), foreground=palette.get("button_fg"))
        self.style.map("Treeview.Heading", background=[('active', palette.get("header_fg"))]) 
        self.style.configure("Treeview", rowheight=25, font=('Helvetica', 10), 
                             background=palette.get("entry_bg"), foreground=palette.get("entry_fg"),
                             fieldbackground=palette.get("entry_bg")) 
        self.style.map("Treeview", 
                       background=[('selected', palette.get("header_fg", "#0078D7"))],
                       foreground=[('selected', palette.get("button_fg", "white"))])

        self.root.option_add('*Listbox.background', palette["entry_bg"])
        self.root.option_add('*Listbox.foreground', palette["entry_fg"])
        self.root.option_add('*Listbox.selectBackground', palette.get("header_fg", "#0078D7"))
        self.root.option_add('*Listbox.selectForeground', palette.get("button_fg", "white"))

        self.root.option_add('*Text.background', palette["entry_bg"])
        self.root.option_add('*Text.foreground', palette["entry_fg"])
        self.root.option_add('*Text.insertBackground', palette["entry_fg"]) 
        self.root.option_add('*Text.selectBackground', palette.get("header_fg", "#0078D7"))
        self.root.option_add('*Text.selectForeground', palette.get("button_fg", "white"))

        if self.main_frame: 
            self.main_frame.destroy()
            self.create_widgets() 
        if hasattr(self, 'status_bar') and self.status_bar:
            self.status_bar.destroy()
            self.create_status_bar() 

    def create_menu(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        menu_bg = self.current_palette.get("frame_bg", self.root.cget('bg')) 
        menu_fg = self.current_palette.get("fg", "black")

        file_menu = tk.Menu(menubar, tearoff=0, bg=menu_bg, fg=menu_fg)
        file_menu.add_command(label="Settings", command=self.open_settings_window)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        menubar.add_cascade(label="File", menu=file_menu)

        view_menu = tk.Menu(menubar, tearoff=0, bg=menu_bg, fg=menu_fg)
        view_menu.add_command(label=self.get_ui_label("view_data_window_title", "View Patient Data"), 
                              command=self.view_patient_data_window)
        view_menu.add_command(label=self.get_ui_label("view_served_worklist_title", "View Served Worklist"),
                              command=self.open_served_worklist_viewer)
        menubar.add_cascade(label="View", menu=view_menu)
        logging.debug("Menu created.")

    def create_fields(self):
        fields_config = [
    ("patient_id", "<FocusOut>", self.on_patient_id_change, "Unique patient ID. Auto-formats (e.g., 123456AB -> 123456 AB). Fetches existing data on FocusOut."),
    ("patient_name", None, None, "Full name of the patient."),
    ("accession", None, None, "Accession Number (e.g., CRHCT123, CRHDX456). Must be unique per study."),
    ("dob", None, None, "Patient's date of birth (DD/MM/YYYY or DDMMYYYY)."),
    ("sex", None, None, "Sex (M for Male, F for Female)."),
    ("study_description", None, None, "Description of the study or examination."),
    ("referred_from", None, None, "Referring clinic or doctor."),
    ("requesting_physician", None, None, "Name of the physician requesting the study.")
]

        
        self.field_widgets = {} 

        for i, (key, bind_event, bind_function, tooltip_text) in enumerate(fields_config):
            label_text = self.get_ui_label(key, key.replace("_", " ").title() + ":")
            label = ttk.Label(self.main_frame, text=label_text)
            label.grid(row=i, column=0, padx=5, pady=7, sticky=tk.W)

            entry_var_name = f"entry_{key}" 

            if key in ["study_description", "referred_from", "requesting_physician", "scheduled_station_ae"]:
                entry_widget = ttk.Combobox(self.main_frame, width=38, font=('Helvetica', 11))
                if key == "scheduled_station_ae": 
                    entry_widget.set(self.config.get("Preferences", "default_scheduled_station_ae", fallback=""))
            else:
                entry_widget = ttk.Entry(self.main_frame, width=40, font=('Helvetica', 11))
            
            entry_widget.grid(row=i, column=1, padx=5, pady=7, sticky=tk.EW)
            setattr(self, entry_var_name, entry_widget) 
            self.field_widgets[entry_var_name] = entry_widget 

            ToolTip(entry_widget, tooltip_text, self.config, self)

            if bind_event and bind_function:
                entry_widget.bind(bind_event, bind_function)
        
        self.num_fields = len(fields_config) 
        logging.debug(f"{self.num_fields} input fields created on main form.")

    def create_buttons(self):
        button_frame_row = self.num_fields + 1 
        button_frame = ttk.Frame(self.main_frame, style='Custom.TFrame')
        button_frame.grid(row=button_frame_row, column=0, columnspan=2, pady=12, sticky=tk.EW)

        button_frame.columnconfigure(0, weight=1) 
        button_frame.columnconfigure(1, weight=1) 
        button_frame.columnconfigure(2, weight=1) 

        self.submit_button = ttk.Button(button_frame, 
                                        text=self.get_ui_label("submit_button", "Register Patient"), 
                                        command=self.submit_form, width=25)
        self.submit_button.grid(row=0, column=0, padx=5, pady=5, sticky=tk.E)
        ToolTip(self.submit_button, "Register the patient with the entered details. Data will be available to the MWL server.", self.config, self)

        self.clear_button = ttk.Button(button_frame, 
                                       text=self.get_ui_label("clear_button", "Clear Form"), 
                                       command=self.confirm_clear_form, width=15)
        self.clear_button.grid(row=0, column=1, padx=5, pady=5) 
        ToolTip(self.clear_button, "Clear all input fields and any manually staged attachments.", self.config, self)
        
        self.email_button = ttk.Button(button_frame,
                                       text=self.get_ui_label("email_button", "Email Previous Report"),
                                       command=self.trigger_email_report_picker, width=20)
        self.email_button.grid(row=0, column=2, padx=5, pady=5, sticky=tk.W)
        ToolTip(self.email_button, "Select a previously generated DOCX report to compose an email with its associated attachments.", self.config, self)
        logging.debug("Action buttons created on main form.")

    def create_widgets(self):
        if self.main_frame and self.main_frame.winfo_exists(): 
            self.main_frame.destroy()

        self.main_frame = ttk.Frame(self.root, padding="15", style='Custom.TFrame')
        self.main_frame.pack(expand=True, fill=tk.BOTH)
        self.main_frame.columnconfigure(1, weight=1) 

        self.create_fields() 

        attach_frame_row = self.num_fields 
        attach_frame = ttk.Frame(self.main_frame, style='Custom.TFrame')
        attach_frame.grid(row=attach_frame_row, column=0, columnspan=2, pady=8, sticky=tk.EW)

        self.attach_button = ttk.Button(attach_frame, 
                                        text=self.get_ui_label("attach_files_button", "Attach File(s)"), 
                                        command=self.select_manual_attachments, width=18)
        self.attach_button.pack(side=tk.LEFT, padx=(0, 5))
        ToolTip(self.attach_button, "Manually select files to be included with the generated DOCX report. Files in watched folders may also be auto-staged here.", self.config, self)
        
        self.manual_attach_label_var = tk.StringVar(
            value=f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}"
        )
        manual_attach_display_label = ttk.Label(attach_frame, textvariable=self.manual_attach_label_var)
        manual_attach_display_label.pack(side=tk.LEFT, padx=(10, 0))

        self.create_buttons() 
        logging.debug("Main frame widgets (fields, attachments, buttons) created/recreated.")

    def select_manual_attachments(self):
        file_types = (
            ("All files", "*.*"),
            ("PDF files", "*.pdf"),
            ("Image files", "*.jpg;*.jpeg;*.png;*.tiff;*.bmp"),
            ("DICOM files", "*.dcm"),
            ("Document files", "*.doc;*.docx;*.txt")
        )
        filepaths = filedialog.askopenfilenames(
            title="Select Files to Attach", 
            parent=self.root, 
            filetypes=file_types
        )
        if filepaths: 
            newly_added_count = 0
            for fp in filepaths:
                if fp not in self.manual_attachments:
                    self.manual_attachments.append(fp)
                    newly_added_count +=1
            
            if newly_added_count > 0:
                self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")
                self.update_status(f"{newly_added_count} file(s) manually staged for attachment.")
                logging.info(f"{newly_added_count} files manually staged. Total attachments: {len(self.manual_attachments)}")

    def create_status_bar(self):
        if hasattr(self, 'status_bar') and self.status_bar and self.status_bar.winfo_exists():
            self.status_bar.destroy()
            
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var, 
                                    relief=tk.SUNKEN, anchor=tk.W, padding=(5,2),
                                    style="Status.TLabel") 
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        logging.debug("Status bar created/recreated.")

    def update_status(self, message, is_error=False, duration=5000):
        if not (hasattr(self, 'status_var') and self.status_var): 
            logging.warning(f"Status bar not ready for message: {message}")
            return

        self.status_var.set(message)
        if is_error:
            logging.error(f"Status Update (Error): {message}")
        else:
            logging.info(f"Status Update: {message}")

        if duration > 0: 
            self.root.after(duration, lambda: self.status_var.set("Ready") if self.status_var.get() == message else None)

    def get_modality_specific_path(self, base_section_key, modality_code):
        if modality_code and self.config.has_option(base_section_key, modality_code):
            path = self.config.get(base_section_key, modality_code, fallback=None)
            if path: return os.path.normpath(os.path.expanduser(path))
        
        if self.config.has_option(base_section_key, "Default"):
            path = self.config.get(base_section_key, "Default", fallback=None)
            if path: return os.path.normpath(os.path.expanduser(path))
            
        if "Output.DOCX" in base_section_key:
            return os.path.normpath(os.path.expanduser(self.config.get("Paths", "general_docx_output_folder", fallback="")))
        elif "WatchFolders" in base_section_key:
             path = self.config.get("Paths", "general_watch_folder", fallback="")
             return os.path.normpath(os.path.expanduser(path)) if path else "" 
        
        logging.warning(f"Could not determine a specific path for section '{base_section_key}' and modality '{modality_code}'.")
        return "" 

    def load_combobox_values_from_db(self):
        try:
            referred_from_values = get_distinct_values_for_combobox_db("referred_from")
            study_desc_values = get_distinct_values_for_combobox_db("study_description")
            req_phys_values = get_distinct_values_for_combobox_db("requesting_physician")
            sched_ae_values = get_distinct_values_for_combobox_db("scheduled_station_ae_title")

            if hasattr(self, 'entry_referred_from'):
                self.entry_referred_from['values'] = sorted(list(referred_from_values))
            if hasattr(self, 'entry_study_description'):
                self.entry_study_description['values'] = sorted(list(study_desc_values))
            if hasattr(self, 'entry_requesting_physician'):
                self.entry_requesting_physician['values'] = sorted(list(req_phys_values))
            if hasattr(self, 'entry_scheduled_station_ae'):
                self.entry_scheduled_station_ae['values'] = sorted(list(sched_ae_values))
                
            logging.info("Combobox values loaded from Database.")
        except Exception as e:
            self.update_status(f"Error loading combobox values from DB: {e}", True)
            logging.exception("Error loading combobox values from Database.")
            
    def check_for_duplicates(self, patient_name, patient_id, accession_number):
        is_dup, hrs, mins, prev_dt, prev_desc, match_type = check_duplicate_record_db(patient_name, patient_id, accession_number)
        if is_dup:
            logging.info(f"Duplicate check: Found recent entry for {patient_name}/{patient_id} (Type: {match_type})")
        return is_dup, hrs, mins, prev_dt, prev_desc 

    def generate_docx_report(self, data_dict, modality_code, patient_specific_base_path):
        template_path = self.config.get("Paths", "docx_template", fallback="")
        if not template_path or not os.path.exists(template_path):
            messagebox.showerror("DOCX Template Error", f"DOCX template not found at:\n{template_path}\nPlease check Settings > General Paths.", parent=self.root)
            self.update_status(f"DOCX template not found: {template_path}", True)
            logging.error(f"DOCX template not found: {template_path}")
            return None
        
        try:
            doc = Document(template_path)
        except Exception as e:
            messagebox.showerror("DOCX Load Error", f"Error loading template '{os.path.basename(template_path)}':\n{e}", parent=self.root)
            self.update_status(f"Error loading DOCX template: {e}", True)
            logging.exception(f"Error loading DOCX template: {template_path}")
            return None

        dob_f = format_date_friendly(data_dict.get("Date of Birth", "")) 
        study_date_f = format_date_friendly(data_dict.get("Study Date", "")) 
        sex_val = data_dict.get("Sex", "")
        sex_f = "Male" if sex_val == "M" else "Female" if sex_val == "F" else sex_val

        # FIXED: Added missing placeholders for template
        replacements = {
            '{Patient Name}': data_dict.get("Patient Name", ""),
            '{Docket Number}': data_dict.get("Patient ID", ""), 
            '{Date of Birth}': dob_f,
            '{Accession Number}': data_dict.get("Accession Number", ""),
            '{Study Description}': data_dict.get("Study Description", ""),
            '{Referring Physician}': data_dict.get("Referred From", ""), 
            '{Requesting Physician}': data_dict.get("Requesting Physician", ""),
            '{Study Date}': study_date_f,
            '{Date of Exam}': study_date_f,  # ADDED: Missing placeholder
            '{Clinic Referred From}': data_dict.get("Referred From", ""),  # ADDED: Missing placeholder
            '{Body Part Done}': data_dict.get("Study Description", ""),  # ADDED: Missing placeholder (same as study description)
            '{Modality}': data_dict.get("Modality", modality_code), 
            '{Modality Done}': data_dict.get("Modality", modality_code),  # ADDED: Missing placeholder
            '{Sex}': sex_f
        }

        for p in doc.paragraphs:
            for r in p.runs:
                for k, v in replacements.items():
                    if k in r.text:
                        r.text = r.text.replace(k, str(v))
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p_cell in cell.paragraphs:
                        for r_cell in p_cell.runs:
                            for k,v in replacements.items():
                                if k in r_cell.text:
                                    r_cell.text = r_cell.text.replace(k,str(v))
        
        safe_name = "".join(c if c.isalnum() else "_" for c in data_dict.get("Patient Name", "UnknownPatient"))
        safe_desc = "".join(c if c.isalnum() else "_" for c in data_dict.get("Study Description", "NoDesc"))[:30] 
        study_date_filename_part = study_date_f.replace(',', '').replace(' ', '_') if study_date_f != "N/A" else datetime.now().strftime("%b_%d_%Y")
        
        fname = f"{safe_name}_{modality_code}_{safe_desc}_{study_date_filename_part}_REPORT.docx"
        output_file_path = os.path.join(patient_specific_base_path, fname)

        try:
            doc.save(output_file_path)
            self.update_status(f"DOCX report generated: {fname}")
            logging.info(f"DOCX report generated: {output_file_path} for Patient ID {data_dict.get('Patient ID')}")
            return output_file_path
        except Exception as e:
            messagebox.showerror("DOCX Save Error", f"Error saving DOCX report '{fname}':\n{e}", parent=self.root)
            self.update_status(f"Error saving DOCX: {e}", True)
            logging.exception(f"Error saving DOCX file {output_file_path} for Patient ID {data_dict.get('Patient ID')}")
            return None

    def process_attachments(self, patient_id_for_match, modality_code, patient_specific_report_folder):
        all_copied_to_report_folder = [] 
        copied_count_manual = 0

        if self.manual_attachments:
            ensure_dir_exists(patient_specific_report_folder) 
            remaining_manual_attachments = [] 
            for src_path in self.manual_attachments:
                if os.path.isfile(src_path):
                    try:
                        filename = os.path.basename(src_path)
                        dest_path = os.path.join(patient_specific_report_folder, filename)
                        shutil.copy2(src_path, dest_path)
                        all_copied_to_report_folder.append(dest_path)
                        copied_count_manual += 1
                        logging.info(f"Copied staged attachment '{filename}' to report folder: {dest_path}")
                    except Exception as e:
                        self.update_status(f"Error copying staged attachment '{filename}': {e}", True)
                        logging.exception(f"Error copying staged attachment '{filename}' from {src_path} to {patient_specific_report_folder}")
                        remaining_manual_attachments.append(src_path) 
                else:
                    logging.warning(f"Staged attachment not found or not a file, removing from list: {src_path}")
            
            self.manual_attachments = remaining_manual_attachments 
            if copied_count_manual > 0:
                self.update_status(f"Processed {copied_count_manual} manually staged attachment(s).")
            self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} {len(self.manual_attachments)}")

        watch_folder = self.get_modality_specific_path("Paths.WatchFolders.Modalities", modality_code)
        if watch_folder and os.path.isdir(watch_folder):
            copied_dcm_count = 0
            try:
                for item_name in os.listdir(watch_folder):
                    if patient_id_for_match.lower() in item_name.lower() and item_name.lower().endswith(".dcm"):
                        src = os.path.join(watch_folder, item_name)
                        dest = os.path.join(patient_specific_report_folder, item_name)
                        
                        if dest not in all_copied_to_report_folder: 
                            try:
                                ensure_dir_exists(patient_specific_report_folder) 
                                shutil.copy2(src, dest)
                                all_copied_to_report_folder.append(dest)
                                copied_dcm_count += 1
                                logging.info(f"Copied DICOM file '{item_name}' from watch folder '{watch_folder}' to report folder.")
                            except Exception as e:
                                self.update_status(f"Error copying DICOM '{item_name}': {e}", True)
                                logging.exception(f"Error copying DICOM file '{item_name}' from {watch_folder}")
                if copied_dcm_count > 0:
                    self.update_status(f"Copied {copied_dcm_count} DICOM file(s) from watch folder for {modality_code}.")
            except Exception as e:
                self.update_status(f"Error accessing watch folder '{watch_folder}' for DICOMs: {e}", True)
                logging.exception(f"Error accessing watch folder '{watch_folder}' for DICOMs")
        
        return all_copied_to_report_folder 

    def submit_form(self):
        logging.info("Submit form initiated.")
        patient_name = self.entry_patient_name.get().strip()
        patient_id_input = self.entry_patient_id.get().strip().upper() 
        accession_number = self.entry_accession.get().strip().upper()
        dob_input_str = self.entry_dob.get().strip()
        sex = self.entry_sex.get().strip().upper() 
        study_description = self.entry_study_description.get().strip()
        referred_from_original_case = self.entry_referred_from.get().strip() 
        requesting_physician = self.entry_requesting_physician.get().strip()
        requested_procedure_id = accession_number  # Auto-generate using accession number
        scheduled_station_ae = self.config.get("Preferences", "default_scheduled_station_ae", fallback="ANY_MODALITY")  # Auto-generate from config

        referred_from_for_check = referred_from_original_case.upper() 

        dob_yyyymmdd = ""; parsed_successfully = False
        if dob_input_str:
            cleaned_dob_input = "".join(filter(str.isdigit, dob_input_str)) 
            if len(cleaned_dob_input) == 8: 
                for fmt_in in ["%d%m%Y", "%Y%m%d", "%m%d%Y"]: 
                    try:
                        dt_obj = datetime.strptime(cleaned_dob_input, fmt_in)
                        dob_yyyymmdd = dt_obj.strftime("%Y%m%d")
                        parsed_successfully = True; break
                    except ValueError: continue
            if not parsed_successfully: 
                for sep_fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y"]:
                    try:
                        dt_obj = datetime.strptime(dob_input_str, sep_fmt)
                        dob_yyyymmdd = dt_obj.strftime("%Y%m%d")
                        parsed_successfully = True; break
                    except ValueError: continue
        
        required_fields_map = {
    "patient_name": patient_name, "patient_id": patient_id_input, "accession": accession_number, 
    "dob": dob_input_str, 
    "sex": sex, "study_description": study_description, 
    "referred_from": referred_from_original_case, 
    "requesting_physician": requesting_physician
}
        for key, val in required_fields_map.items():
            if not val:
                field_display_name = self.get_ui_label(key, key.replace("_", " ").title())
                messagebox.showerror("Validation Error", f"{field_display_name} is required!", parent=self.root)
                self.update_status(f"Validation Error: {field_display_name} required.", True)
                logging.warning(f"Validation Error: {field_display_name} is required.")
                return
        
        if not dob_yyyymmdd: 
            messagebox.showerror("Validation Error", f"Date of Birth '{dob_input_str}' is invalid. Please use DD/MM/YYYY or DDMMYYYY.", parent=self.root)
            self.update_status("Invalid Date of Birth format.", True)
            logging.warning(f"Invalid DOB entered: {dob_input_str}")
            return

        modality = get_modality_from_accession(accession_number)
        if not modality:
            messagebox.showerror("Validation Error", "Invalid Accession Number format. Must start with a known modality prefix (e.g., CRHCT, CRHDX).", parent=self.root)
            self.update_status("Invalid Accession prefix.", True)
            logging.warning(f"Invalid Accession prefix: {accession_number}")
            return
            
        processed_pid = patient_id_input
        if ' ' not in patient_id_input and len(patient_id_input) >= 8 and patient_id_input[:6].isdigit() and patient_id_input[6:].isalnum():
            processed_pid = f"{patient_id_input[:6]} {patient_id_input[6:]}"
        elif len(processed_pid) < 7: 
            messagebox.showerror("Validation Error", "Patient ID is too short or has an invalid format.", parent=self.root)
            self.update_status("Invalid Patient ID format.", True)
            logging.warning(f"Invalid Patient ID format: {patient_id_input}")
            return

        if sex not in ['M', 'F', 'O']: 
            messagebox.showerror("Validation Error", "Sex must be 'M', 'F', or 'O'!", parent=self.root)
            self.update_status("Invalid Sex value.", True)
            logging.warning(f"Invalid Sex entered: {sex}")
            return

        # FIXED: Improved duplicate handling to allow continuation after user confirmation
        is_dup, hrs, mins, prev_dt_str, prev_desc = self.check_for_duplicates(patient_name, processed_pid, accession_number)
        user_confirmed_duplicate = False
        if is_dup:
            dup_msg = f"A recent record for Patient '{patient_name}' (or ID '{processed_pid}') was found.\n"
            dup_msg += f"Previous Study: {prev_desc}\nRegistered: ~{hrs}h {mins}m ago ({prev_dt_str}).\n\n"
            dup_msg += "Do you want to register this new study anyway?"
            user_confirmed_duplicate = messagebox.askyesno("Duplicate Warning", dup_msg, parent=self.root)
            if not user_confirmed_duplicate:
                self.update_status("Registration cancelled by user (duplicate).")
                logging.info(f"Registration cancelled due to duplicate for {patient_name}/{processed_pid}")
                self.clear_form_fields() 
                return
            logging.info(f"Duplicate warning for {patient_name}/{processed_pid} overridden by user.")

        study_date_now = datetime.now().strftime("%Y%m%d")
        study_time_now = datetime.now().strftime("%H%M%S")

        patient_data = {
    "Patient Name": patient_name, "Patient ID": processed_pid, 
    "Accession Number": accession_number, "Date of Birth": dob_yyyymmdd,
    "Sex": sex, "Study Date": study_date_now, "Study Time": study_time_now,
    "Study Description": study_description, "Referred From": referred_from_original_case,
    "Modality": modality, "Requesting Physician": requesting_physician,
    "Requested Procedure ID": requested_procedure_id, "Scheduled Station AE Title": scheduled_station_ae
}

        try:
            patient_record_id = add_patient_record_db(patient_data)
            if patient_record_id:
                self.update_status("Patient data saved. Available to MWL Server.")
                self.load_combobox_values_from_db() 
                logging.info(f"Patient data saved to DB for PID: {processed_pid}, Accession: {accession_number}, RecordID: {patient_record_id}. Data available to MWL SCP.")
            else:
                messagebox.showerror("Database Error", "Failed to save patient data to database. Check logs for details.", parent=self.root)
                self.update_status("Error saving patient data to DB.", True)
                return 
        except sqlite3.IntegrityError as e:
             messagebox.showerror("Database Integrity Error", f"Could not save data: {e}.\nPatient ID or Accession Number might already exist with different conflicting data or a database constraint was violated.", parent=self.root)
             self.update_status(f"DB integrity error: {e}", True)
             logging.error(f"DB IntegrityError for {processed_pid}/{accession_number}: {e}")
             return
        except Exception as e: 
            messagebox.showerror("Database Error", f"An unexpected error occurred while saving to the database: {e}", parent=self.root)
            self.update_status(f"Unexpected DB error: {e}", True)
            logging.exception(f"Failed to save data to DB for Patient ID {processed_pid}")
            return

        modality_base_output_folder = self.get_modality_specific_path("Paths.Output.DOCX.Modalities", modality)
        if not modality_base_output_folder:
            messagebox.showerror("Configuration Error", f"Base output path for modality '{modality}' reports is not configured!\nPlease check Settings > Modality Paths.", parent=self.root)
            self.update_status(f"Base output path for {modality} missing.", True)
            logging.error(f"Base output path for modality {modality} not configured.")

        safe_pname_folder = "".join(c if c.isalnum() else "_" for c in patient_name)
        safe_pid_folder = processed_pid.replace(' ', '_') 
        patient_subfolder_name = f"{safe_pname_folder}_{safe_pid_folder}" 
        final_patient_report_folder = os.path.join(modality_base_output_folder, patient_subfolder_name)

        generate_report_and_folder = True
        if modality == "DX": 
            referred_ok = any(referred_from_for_check.startswith(p) for p in ["H/", "HE", "PP", "PR"])
            if not (referred_ok or len(self.manual_attachments) > 0):
                generate_report_and_folder = False
                self.update_status(f"DOCX for {modality} skipped: Criteria not met (Referral/Attachment).", duration=0) 
                logging.info(f"DOCX report generation for {modality} patient {processed_pid} skipped: Referral prefix or manual attachment criteria not met.")
                if self.manual_attachments: 
                    self.manual_attachments.clear()
                    self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} 0")
                    logging.info("Cleared manual attachments as DX report was skipped.")
                messagebox.showinfo("Partial Registration (DX)", 
                                    f"Patient '{patient_name}' registered (data available to MWL Server).\n\n"
                                    f"DOCX Report for {modality} was NOT generated due to missing referral prefix or manual attachment.",
                                    parent=self.root)
        
        generated_docx_path = None
        if generate_report_and_folder and modality_base_output_folder: 
            ensure_dir_exists(final_patient_report_folder)
            generated_docx_path = self.generate_docx_report(patient_data, modality, final_patient_report_folder)
            if generated_docx_path:
                self.process_attachments(processed_pid, modality, final_patient_report_folder)
            else:
                logging.error(f"DOCX generation failed for {processed_pid}, but generate_report_and_folder was True.")
                if self.manual_attachments: 
                     self.update_status("DOCX generation failed; attachments remain staged.", 3000)

        self.config.set("Preferences", "last_referred_from", referred_from_original_case)
        save_config(self.config)

        if generate_report_and_folder and not generated_docx_path and modality_base_output_folder:
            self.update_status("Patient registered. DOCX report generation failed.", 0) 
        elif generate_report_and_folder and generated_docx_path:
            messagebox.showinfo("Success", 
                                f"PATIENT REGISTERED!\n\nPatient: {patient_name} ({processed_pid})\n"
                                f"Report & attachments processed to:\n{final_patient_report_folder}\n\n"
                                "Data available to MWL Server.", 
                                parent=self.root)
            self.update_status("Patient registered. Report & attachments processed.")
            logging.info(f"Successful full registration for {processed_pid}. Report: {generated_docx_path}")
        elif not generate_report_and_folder: 
            logging.info(f"Partial registration for {processed_pid} (DX, no report generated).")
        
        self.clear_form_fields()
        if hasattr(self, 'entry_patient_id'): 
            self.entry_patient_id.focus()
        logging.info("Form submission process completed.")

    def confirm_clear_form(self):
        if messagebox.askyesno("Confirm Clear", "Are you sure you want to clear all fields and staged attachments?", parent=self.root):
            self.clear_form_fields()
            self.update_status("Form cleared.")
            logging.info("Form cleared by user confirmation.")

    def clear_form_fields(self):
     for attr_name in ['entry_patient_id', 'entry_patient_name', 'entry_dob', 'entry_sex']: 
        if hasattr(self, attr_name):
            getattr(self, attr_name).delete(0, tk.END)
    
     if hasattr(self, 'entry_accession'):
        self.entry_accession.delete(0, tk.END)
        self.entry_accession.insert(0, self.config.get("Preferences", "default_accession_prefix", fallback="CRH"))

     for combo_attr in ['entry_study_description', 'entry_referred_from', 'entry_requesting_physician']:
        if hasattr(self, combo_attr) and isinstance(getattr(self, combo_attr), ttk.Combobox):
            getattr(self, combo_attr).set('') 
    
     self.manual_attachments.clear()
     self.manual_attach_label_var.set(f"{self.get_ui_label('attachments_label_prefix', 'Attachments:')} 0")
    
     logging.info("Form fields and manual attachments cleared.")
     if hasattr(self, 'entry_patient_id'): 
        self.entry_patient_id.focus()

    def on_patient_id_change(self, event=None): 
        if not hasattr(self, 'entry_patient_id'): return 

        pid_in = self.entry_patient_id.get().strip().upper()
        if not pid_in: return 

        current_cursor_pos = self.entry_patient_id.index(tk.INSERT)
        formatted_pid = pid_in

        if ' ' not in pid_in and len(pid_in) >= 8 and \
           pid_in[:6].isdigit() and pid_in[6:].isalnum(): 
            formatted_pid = f"{pid_in[:6]} {pid_in[6:]}"
            self.entry_patient_id.delete(0, tk.END)
            self.entry_patient_id.insert(0, formatted_pid)
            try:
                self.entry_patient_id.icursor(current_cursor_pos + 1 if current_cursor_pos >= 6 else current_cursor_pos)
            except tk.TclError: 
                pass 
        
        try: self.entry_patient_id.unbind("<FocusOut>")
        except tk.TclError: pass 

        data = get_patient_by_id_db(formatted_pid)
        if data:
            self.populate_fields(data) 
            self.update_status(f"Data loaded from DB for Patient ID: {formatted_pid}")
            logging.info(f"Patient data loaded from DB for ID: {formatted_pid}")
        else:
            for attr in ['entry_patient_name', 'entry_dob', 'entry_sex', 
                         'entry_requesting_physician', 'entry_requested_procedure_id', 
                         'entry_scheduled_station_ae']: 
                if hasattr(self, attr):
                    widget = getattr(self,attr)
                    if isinstance(widget, ttk.Combobox): widget.set("")
                    else: widget.delete(0, tk.END)
            self.update_status(f"No data found in DB for Patient ID: {formatted_pid}")
            logging.info(f"No patient data found for ID: {formatted_pid}")

        try:
            self.root.after(100, lambda: self.entry_patient_id.bind("<FocusOut>", self.on_patient_id_change) if hasattr(self, 'entry_patient_id') and self.entry_patient_id.winfo_exists() else None)
        except tk.TclError: pass 

    def populate_fields(self, patient_data):
        if hasattr(self, 'entry_patient_name'):
            self.entry_patient_name.delete(0, tk.END)
            self.entry_patient_name.insert(0, patient_data.get('patient_name', ''))
        
        if hasattr(self, 'entry_dob'):
            self.entry_dob.delete(0, tk.END)
            dob_yyyymmdd = patient_data.get('dob_yyyymmdd', '')
            if dob_yyyymmdd:
                try:
                    dob_display = datetime.strptime(dob_yyyymmdd, "%Y%m%d").strftime("%d/%m/%Y")
                    self.entry_dob.insert(0, dob_display)
                except ValueError:
                    self.entry_dob.insert(0, dob_yyyymmdd) 
        
        if hasattr(self, 'entry_sex'):
            self.entry_sex.delete(0, tk.END)
            self.entry_sex.insert(0, patient_data.get('sex', ''))

        if patient_data.get('referred_from') and hasattr(self, 'entry_referred_from') and isinstance(self.entry_referred_from, ttk.Combobox):
            self.entry_referred_from.set(patient_data.get('referred_from', ''))
        
        if patient_data.get('requesting_physician') and hasattr(self, 'entry_requesting_physician') and isinstance(self.entry_requesting_physician, ttk.Combobox):
            self.entry_requesting_physician.set(patient_data.get('requesting_physician', ''))

        if patient_data.get('scheduled_station_ae_title') and hasattr(self, 'entry_scheduled_station_ae') and isinstance(self.entry_scheduled_station_ae, ttk.Combobox):
            self.entry_scheduled_station_ae.set(patient_data.get('scheduled_station_ae_title', ''))
        
        
            
        if hasattr(self, 'entry_study_description') and isinstance(self.entry_study_description, ttk.Combobox):
            self.entry_study_description.set('') 
        if hasattr(self, 'entry_accession'):
            self.entry_accession.delete(0, tk.END)
            self.entry_accession.insert(0, self.config.get("Preferences", "default_accession_prefix", fallback="CRH"))
        
        logging.debug(f"Populated form fields with existing data for Patient ID {patient_data.get('patient_id')}")

    def open_settings_window(self):
        logging.info("Opening settings window.")
        settings_win = tk.Toplevel(self.root)
        settings_win.title(self.get_ui_label("settings_window_title", "Settings"))
        settings_win.geometry("950x850") 
        settings_win.transient(self.root) 
        settings_win.grab_set() 
        settings_win.configure(bg=self.current_palette.get("bg", "#F0F0F0"))

        tab_control = ttk.Notebook(settings_win)

        paths_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        modality_paths_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        app_mwl_server_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        ext_reports_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        appearance_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        prefs_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        smtp_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        email_recipients_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        email_templates_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)
        ui_labels_tab = ttk.Frame(tab_control, style='Custom.TFrame', padding=10)

        tab_control.add(paths_tab, text='General Paths')
        tab_control.add(modality_paths_tab, text='Modality Paths')
        tab_control.add(app_mwl_server_tab, text=self.get_ui_label("app_mwl_server_tab_title", "This App as MWL Server"))
        tab_control.add(ext_reports_tab, text=self.get_ui_label("ext_reports_tab_title", "External Reports"))
        tab_control.add(appearance_tab, text=self.get_ui_label("appearance_tab_title", "Appearance"))
        tab_control.add(prefs_tab, text='Preferences')
        tab_control.add(smtp_tab, text=self.get_ui_label("smtp_settings_tab_title", "SMTP Email"))
        tab_control.add(email_recipients_tab, text=self.get_ui_label("email_recipients_tab_title", "Email Recipients"))
        tab_control.add(email_templates_tab, text=self.get_ui_label("email_templates_tab_title", "Email Templates"))
        tab_control.add(ui_labels_tab, text='UI Labels')
        
        tab_control.pack(expand=1, fill="both", padx=10, pady=10)

        self.settings_entries = {} 
        self.ui_label_settings_entries = {}
        self.ext_report_watch_folders_listbox = None 
        self.favorite_recipients_listbox = None
        self.email_templates_listbox = None
        self.current_template_widgets = {} 
        self.app_mwl_config_widgets = {} 

        ttk.Label(paths_tab, text="General File Paths:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        gp_frame = ttk.Frame(paths_tab, style='Custom.TFrame')
        gp_frame.pack(expand=True, fill=tk.BOTH)
        gp_frame.columnconfigure(1, weight=1) 
        
        general_paths_map = [
            ("Database File:", "Paths", "db_file", False), 
            ("DOCX Template File:", "Paths", "docx_template", False),
            ("General DOCX Output Folder:", "Paths", "general_docx_output_folder", True), 
            ("General Modality Watch Folder (optional):", "Paths", "general_watch_folder", True)
        ]
        for r, (lbl_text, section, key, is_folder) in enumerate(general_paths_map):
            ttk.Label(gp_frame, text=lbl_text).grid(row=r, column=0, sticky=tk.W, padx=5, pady=7)
            entry = ttk.Entry(gp_frame, width=70)
            entry.insert(0, self.config.get(section, key, fallback=""))
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=7)
            if key != "db_file": 
                ttk.Button(gp_frame, text="Browse...", 
                           command=lambda e=entry, f=is_folder: self.browse_path(e, f, parent=settings_win)
                          ).grid(row=r, column=2, padx=5, pady=7)
            else:
                entry.config(state="readonly") 
            self.settings_entries[(section, key)] = entry

        ttk.Label(modality_paths_tab, text="Modality-Specific Paths:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        mp_canvas = tk.Canvas(modality_paths_tab, bg=self.current_palette.get("frame_bg"), highlightthickness=0)
        mp_scrollbar = ttk.Scrollbar(modality_paths_tab, orient="vertical", command=mp_canvas.yview)
        mp_scrollable_frame = ttk.Frame(mp_canvas, style='Custom.TFrame')
        mp_scrollable_frame.bind("<Configure>", lambda e: mp_canvas.configure(scrollregion=mp_canvas.bbox("all")))
        mp_canvas_window = mp_canvas.create_window((0, 0), window=mp_scrollable_frame, anchor="nw")
        mp_canvas.configure(yscrollcommand=mp_scrollbar.set)
        mp_canvas.pack(side="left", fill="both", expand=True)
        mp_scrollbar.pack(side="right", fill="y")
        mp_scrollable_frame.columnconfigure(1, weight=1) 
        mp_scrollable_frame.columnconfigure(4, weight=1)
        row_idx = 0
        for mod_code in MODALITIES:
            ttk.Label(mp_scrollable_frame, text=f"{mod_code} DOCX Output:", font=('Helvetica', 10, 'bold')).grid(row=row_idx, column=0, sticky=tk.W, padx=5, pady=3)
            entry_docx = ttk.Entry(mp_scrollable_frame, width=35)
            entry_docx.insert(0, self.config.get("Paths.Output.DOCX.Modalities", mod_code, fallback=""))
            entry_docx.grid(row=row_idx, column=1, sticky=tk.EW, padx=5, pady=3)
            ttk.Button(mp_scrollable_frame, text="...", width=3, command=lambda e=entry_docx: self.browse_path(e, True, parent=settings_win)).grid(row=row_idx, column=2, padx=(0,10), pady=3)
            self.settings_entries[("Paths.Output.DOCX.Modalities", mod_code)] = entry_docx

            ttk.Label(mp_scrollable_frame, text=f"{mod_code} Watch Folder:", font=('Helvetica', 10, 'bold')).grid(row=row_idx, column=3, sticky=tk.W, padx=(10,5), pady=3)
            entry_watch = ttk.Entry(mp_scrollable_frame, width=35)
            entry_watch.insert(0, self.config.get("Paths.WatchFolders.Modalities", mod_code, fallback=""))
            entry_watch.grid(row=row_idx, column=4, sticky=tk.EW, padx=5, pady=3)
            ttk.Button(mp_scrollable_frame, text="...", width=3, command=lambda e=entry_watch: self.browse_path(e, True, parent=settings_win)).grid(row=row_idx, column=5, padx=(0,5), pady=3)
            self.settings_entries[("Paths.WatchFolders.Modalities", mod_code)] = entry_watch
            row_idx += 1
        mp_scrollable_frame.bind("<Configure>", lambda e: mp_canvas.itemconfig(mp_canvas_window, width=e.width))

        self._setup_app_mwl_server_tab(app_mwl_server_tab)

        ttk.Label(ext_reports_tab, text=self.get_ui_label("ext_reports_watch_folder_label"), font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        er_top_frame = ttk.Frame(ext_reports_tab, style='Custom.TFrame')
        er_top_frame.pack(expand=False, fill=tk.X, pady=(0,10))
        
        self.ext_report_watch_folders_listbox = tk.Listbox(er_top_frame, selectmode=tk.SINGLE, height=6, exportselection=False,
                                                            bg=self.current_palette.get("entry_bg"), 
                                                            fg=self.current_palette.get("entry_fg"),
                                                            selectbackground=self.current_palette.get("header_fg"),
                                                            selectforeground=self.current_palette.get("button_fg"))
        self.ext_report_watch_folders_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, padx=(0,5))
        
        # FIXED: Bind folder selection to load auto-send configuration
        self.ext_report_watch_folders_listbox.bind("<<ListboxSelect>>", self._on_ext_report_folder_select)
        
        er_buttons_frame = ttk.Frame(er_top_frame, style='Custom.TFrame')
        er_buttons_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        ttk.Button(er_buttons_frame, text="Add Folder", command=lambda: self._add_ext_report_folder(settings_win)).pack(pady=(0,5), fill=tk.X)
        ttk.Button(er_buttons_frame, text="Remove Selected", command=self._remove_ext_report_folder).pack(fill=tk.X)

        # FIXED: Auto-send configuration section for external reports
        autosend_frame = ttk.LabelFrame(ext_reports_tab, text=self.get_ui_label("ext_report_autosend_group_label", "Automatic Emailing for Selected Folder:"), 
                                       style='Custom.TFrame', padding=10)
        autosend_frame.pack(expand=False, fill=tk.X, pady=(10,0))
        
        # Enable auto-send checkbox
        self.ext_report_autosend_widgets['enable_var'] = tk.BooleanVar()
        enable_check = ttk.Checkbutton(autosend_frame, text=self.get_ui_label("ext_report_autosend_enable_label", "Enable Auto-Emailing for this Folder"),
                                      variable=self.ext_report_autosend_widgets['enable_var'])
        enable_check.pack(anchor=tk.W, pady=(0,5))
        self.ext_report_autosend_widgets['enable'] = enable_check
        
        # Recipients entry
        recipients_frame = ttk.Frame(autosend_frame, style='Custom.TFrame')
        recipients_frame.pack(fill=tk.X, pady=(0,5))
        ttk.Label(recipients_frame, text=self.get_ui_label("ext_report_autosend_recipients_label", "Auto-Send Recipient(s):")).pack(side=tk.LEFT)
        self.ext_report_autosend_widgets['recipients'] = ttk.Entry(recipients_frame, width=50)
        self.ext_report_autosend_widgets['recipients'].pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(5,0))
        ttk.Button(recipients_frame, text=self.get_ui_label("ext_report_autosend_add_favorite_button", "Add Favorite"), 
                  command=lambda: self._add_favorite_to_autosend_recipients(settings_win)).pack(side=tk.RIGHT, padx=(5,0))
        
        # Template selection
        template_frame = ttk.Frame(autosend_frame, style='Custom.TFrame')
        template_frame.pack(fill=tk.X, pady=(0,5))
        ttk.Label(template_frame, text=self.get_ui_label("ext_report_autosend_template_label", "Auto-Send Email Template:")).pack(side=tk.LEFT)
        self.ext_report_autosend_widgets['template'] = ttk.Combobox(template_frame, width=30, state="readonly")
        self.ext_report_autosend_widgets['template'].pack(side=tk.LEFT, padx=(5,0))
        self.ext_report_autosend_widgets['template'].bind("<<ComboboxSelected>>", self._toggle_autosend_custom_fields_state)
        
        # Custom subject and body (shown when template is 'Custom')
        self.ext_report_autosend_widgets['custom_subject_label'] = ttk.Label(autosend_frame, text=self.get_ui_label("ext_report_autosend_custom_subject_label", "Custom Subject:"))
        self.ext_report_autosend_widgets['custom_subject'] = ttk.Entry(autosend_frame, width=80)
        
        self.ext_report_autosend_widgets['custom_body_label'] = ttk.Label(autosend_frame, text=self.get_ui_label("ext_report_autosend_custom_body_label", "Custom Body:"))
        self.ext_report_autosend_widgets['custom_body'] = tk.Text(autosend_frame, height=4, width=80, 
                                                                  bg=self.current_palette.get("entry_bg"),
                                                                  fg=self.current_palette.get("entry_fg"),
                                                                  insertbackground=self.current_palette.get("entry_fg"))
        
        # Placeholders info
        ttk.Label(autosend_frame, text=self.get_ui_label("ext_report_autosend_placeholders_label", "Placeholders: {Filename}, {FolderPath}, {DateTime}, {Patient Name}, {Modality}, {Study Description}, {Date}"), 
                 foreground="gray").pack(anchor=tk.W, pady=(5,0))

        # Populate external report folders
        for folder in self.get_all_external_report_watch_folders():
            self.ext_report_watch_folders_listbox.insert(tk.END, folder)

        # Appearance tab
        ttk.Label(appearance_tab, text="Visual Appearance Settings:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        app_frame = ttk.Frame(appearance_tab, style='Custom.TFrame')
        app_frame.pack(expand=True, fill=tk.BOTH)
        
        ui_style_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        ui_style_frame.pack(fill=tk.X, pady=(0,10))
        ttk.Label(ui_style_frame, text=self.get_ui_label("ui_style_engine_label", "UI Style Engine:")).pack(side=tk.LEFT)
        ui_style_combo = ttk.Combobox(ui_style_frame, width=25, state="readonly",
                                     values=["System Default", "Clam (Modern)", "Alt (Modern-ish)", "Default (Classic)", "Classic (Older)"])
        ui_style_combo.set(self.config.get("Preferences", "ui_style", fallback="System Default"))
        ui_style_combo.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("Preferences", "ui_style")] = ui_style_combo

        color_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        color_frame.pack(fill=tk.X, pady=(0,10))
        ttk.Label(color_frame, text=self.get_ui_label("color_palette_label", "Color Palette:")).pack(side=tk.LEFT)
        color_combo = ttk.Combobox(color_frame, width=25, state="readonly",
                                  values=["Default", "Light Blue", "Dark", "High Contrast", "Mint Green", "Lavender"])
        color_combo.set(self.config.get("Preferences", "color_theme", fallback="Default"))
        color_combo.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("Preferences", "color_theme")] = color_combo

        ui_size_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        ui_size_frame.pack(fill=tk.X, pady=(0,10))
        ttk.Label(ui_size_frame, text="UI Size:").pack(side=tk.LEFT)
        ui_size_combo = ttk.Combobox(ui_size_frame, width=25, state="readonly",
                                    values=["Very Compact", "Compact", "Default", "Large", "Extra Large"])
        ui_size_combo.set(self.config.get("Preferences", "ui_size", fallback="Default"))
        ui_size_combo.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("Preferences", "ui_size")] = ui_size_combo

        tooltip_frame = ttk.Frame(app_frame, style='Custom.TFrame')
        tooltip_frame.pack(fill=tk.X, pady=(0,10))
        tooltip_var = tk.BooleanVar(value=self.config.getboolean("Preferences", "enable_tooltips", fallback=True))
        ttk.Checkbutton(tooltip_frame, text="Enable Tooltips", variable=tooltip_var).pack(side=tk.LEFT)
        self.settings_entries[("Preferences", "enable_tooltips")] = tooltip_var

        # Preferences tab
        ttk.Label(prefs_tab, text="Application Preferences:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        pref_frame = ttk.Frame(prefs_tab, style='Custom.TFrame')
        pref_frame.pack(expand=True, fill=tk.BOTH)
        pref_frame.columnconfigure(1, weight=1)
        
        prefs_map = [
            ("Default Accession Prefix:", "Preferences", "default_accession_prefix"),
            ("Default Scheduled Station AE:", "Preferences", "default_scheduled_station_ae")
        ]
        for r, (lbl_text, section, key) in enumerate(prefs_map):
            ttk.Label(pref_frame, text=lbl_text).grid(row=r, column=0, sticky=tk.W, padx=5, pady=7)
            entry = ttk.Entry(pref_frame, width=40)
            entry.insert(0, self.config.get(section, key, fallback=""))
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=7)
            self.settings_entries[(section, key)] = entry

        # SMTP tab
        ttk.Label(smtp_tab, text="SMTP Email Configuration:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        smtp_frame = ttk.Frame(smtp_tab, style='Custom.TFrame')
        smtp_frame.pack(expand=True, fill=tk.BOTH)
        smtp_frame.columnconfigure(1, weight=1)
        
        smtp_map = [
            (self.get_ui_label("smtp_server", "SMTP Server:"), "SMTP", "server"),
            (self.get_ui_label("smtp_port", "SMTP Port:"), "SMTP", "port"),
            (self.get_ui_label("smtp_user", "SMTP Username:"), "SMTP", "user"),
            (self.get_ui_label("smtp_password", "SMTP Password:"), "SMTP", "password"),
            (self.get_ui_label("smtp_sender_email", "Sender Email:"), "SMTP", "sender_email")
        ]
        for r, (lbl_text, section, key) in enumerate(smtp_map):
            ttk.Label(smtp_frame, text=lbl_text).grid(row=r, column=0, sticky=tk.W, padx=5, pady=7)
            entry = ttk.Entry(smtp_frame, width=50, show="*" if "password" in key else "")
            entry.insert(0, self.config.get(section, key, fallback=""))
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=7)
            self.settings_entries[(section, key)] = entry

        tls_var = tk.BooleanVar(value=self.config.getboolean("SMTP", "use_tls", fallback=True))
        ttk.Checkbutton(smtp_frame, text=self.get_ui_label("smtp_use_tls", "Use TLS/STARTTLS"), variable=tls_var).grid(row=len(smtp_map), column=0, columnspan=2, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("SMTP", "use_tls")] = tls_var

        ttk.Button(smtp_frame, text=self.get_ui_label("smtp_test_button", "Test SMTP Settings"), 
                  command=lambda: self.test_smtp_settings(settings_win)).grid(row=len(smtp_map)+1, column=0, columnspan=2, pady=10)

        # Email Recipients tab
        ttk.Label(email_recipients_tab, text=self.get_ui_label("favorite_recipients_label", "Favorite Email Recipients:"), 
                 font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        
        rec_frame = ttk.Frame(email_recipients_tab, style='Custom.TFrame')
        rec_frame.pack(expand=True, fill=tk.BOTH)
        
        self.favorite_recipients_listbox = tk.Listbox(rec_frame, selectmode=tk.SINGLE, height=10,
                                                     bg=self.current_palette.get("entry_bg"),
                                                     fg=self.current_palette.get("entry_fg"),
                                                     selectbackground=self.current_palette.get("header_fg"),
                                                     selectforeground=self.current_palette.get("button_fg"))
        self.favorite_recipients_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH, padx=(0,5))
        
        rec_buttons_frame = ttk.Frame(rec_frame, style='Custom.TFrame')
        rec_buttons_frame.pack(side=tk.RIGHT, fill=tk.Y)
        ttk.Button(rec_buttons_frame, text=self.get_ui_label("add_favorite_button", "Add Favorite"), 
                  command=lambda: self._add_favorite_recipient(settings_win)).pack(pady=(0,5), fill=tk.X)
        ttk.Button(rec_buttons_frame, text=self.get_ui_label("remove_favorite_button", "Remove Selected"), 
                  command=self._remove_favorite_recipient).pack(fill=tk.X)
        
        max_recent_frame = ttk.Frame(email_recipients_tab, style='Custom.TFrame')
        max_recent_frame.pack(fill=tk.X, pady=(10,0))
        ttk.Label(max_recent_frame, text=self.get_ui_label("max_recent_recipients_label", "Max Recent Recipients:")).pack(side=tk.LEFT)
        max_recent_entry = ttk.Entry(max_recent_frame, width=10)
        max_recent_entry.insert(0, self.config.get("EmailRecipients", "max_recent", fallback="10"))
        max_recent_entry.pack(side=tk.LEFT, padx=(10,0))
        self.settings_entries[("EmailRecipients", "max_recent")] = max_recent_entry

        # Populate favorite recipients
        favorites_str = self.config.get("EmailRecipients", "favorite_list", fallback="")
        for email in [e.strip() for e in favorites_str.split(';') if e.strip()]:
            self.favorite_recipients_listbox.insert(tk.END, email)

        # Email Templates tab
        ttk.Label(email_templates_tab, text=self.get_ui_label("manage_email_templates_label", "Manage Email Templates:"), 
                 font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        
        templates_main_frame = ttk.Frame(email_templates_tab, style='Custom.TFrame')
        templates_main_frame.pack(expand=True, fill=tk.BOTH)
        
        templates_list_frame = ttk.Frame(templates_main_frame, style='Custom.TFrame')
        templates_list_frame.pack(side=tk.LEFT, fill=tk.Y, padx=(0,10))
        
        self.email_templates_listbox = tk.Listbox(templates_list_frame, selectmode=tk.SINGLE, width=25, height=8,
                                                 bg=self.current_palette.get("entry_bg"),
                                                 fg=self.current_palette.get("entry_fg"),
                                                 selectbackground=self.current_palette.get("header_fg"),
                                                 selectforeground=self.current_palette.get("button_fg"))
        self.email_templates_listbox.pack(expand=True, fill=tk.BOTH, pady=(0,5))
        self.email_templates_listbox.bind("<<ListboxSelect>>", lambda e: self._load_template_for_editing())
        
        temp_buttons_frame = ttk.Frame(templates_list_frame, style='Custom.TFrame')
        temp_buttons_frame.pack(fill=tk.X)
        ttk.Button(temp_buttons_frame, text=self.get_ui_label("add_template_button", "Add New"), 
                  command=lambda: self._edit_email_template(settings_win, is_new=True)).pack(fill=tk.X, pady=(0,2))
        ttk.Button(temp_buttons_frame, text=self.get_ui_label("edit_template_button", "Edit Selected"), 
                  command=lambda: self._edit_email_template(settings_win, is_new=False)).pack(fill=tk.X, pady=(0,2))
        ttk.Button(temp_buttons_frame, text=self.get_ui_label("delete_template_button", "Delete Selected"), 
                  command=self._delete_email_template).pack(fill=tk.X)
        
        templates_edit_frame = ttk.Frame(templates_main_frame, style='Custom.TFrame')
        templates_edit_frame.pack(side=tk.RIGHT, expand=True, fill=tk.BOTH)
        
        ttk.Label(templates_edit_frame, text=self.get_ui_label("template_name_label", "Template Name:")).pack(anchor=tk.W)
        self.current_template_widgets['name'] = ttk.Entry(templates_edit_frame, width=40, state="readonly")
        self.current_template_widgets['name'].pack(fill=tk.X, pady=(0,5))
        
        ttk.Label(templates_edit_frame, text=self.get_ui_label("template_subject_label", "Subject Template:")).pack(anchor=tk.W)
        self.current_template_widgets['subject'] = ttk.Entry(templates_edit_frame, width=80, state="readonly")
        self.current_template_widgets['subject'].pack(fill=tk.X, pady=(0,5))
        
        ttk.Label(templates_edit_frame, text=self.get_ui_label("template_body_label", "Body Template:")).pack(anchor=tk.W)
        self.current_template_widgets['body'] = tk.Text(templates_edit_frame, height=12, width=80, state="disabled",
                                                       bg=self.current_palette.get("entry_bg"),
                                                       fg=self.current_palette.get("entry_fg"))
        self.current_template_widgets['body'].pack(expand=True, fill=tk.BOTH, pady=(0,5))
        
        temp_edit_buttons_frame = ttk.Frame(templates_edit_frame, style='Custom.TFrame')
        temp_edit_buttons_frame.pack(fill=tk.X)
        ttk.Button(temp_edit_buttons_frame, text="Save Changes", command=self._save_current_template_changes).pack(side=tk.LEFT, padx=(0,5))
        ttk.Button(temp_edit_buttons_frame, text="Clear", command=self._clear_template_editor_fields).pack(side=tk.LEFT)
        
        ttk.Label(templates_edit_frame, text=self.get_ui_label("available_placeholders_label", "Placeholders: {Patient Name}, {Modality}, etc."), 
                 foreground="gray").pack(anchor=tk.W, pady=(5,0))

        self._populate_email_templates_listbox()

        # UI Labels tab
        ttk.Label(ui_labels_tab, text="UI Label Customization:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        
        ui_canvas = tk.Canvas(ui_labels_tab, bg=self.current_palette.get("frame_bg"), highlightthickness=0)
        ui_scrollbar = ttk.Scrollbar(ui_labels_tab, orient="vertical", command=ui_canvas.yview)
        ui_scrollable_frame = ttk.Frame(ui_canvas, style='Custom.TFrame')
        ui_scrollable_frame.bind("<Configure>", lambda e: ui_canvas.configure(scrollregion=ui_canvas.bbox("all")))
        ui_canvas_window = ui_canvas.create_window((0, 0), window=ui_scrollable_frame, anchor="nw")
        ui_canvas.configure(yscrollcommand=ui_scrollbar.set)
        ui_canvas.pack(side="left", fill="both", expand=True)
        ui_scrollbar.pack(side="right", fill="y")
        ui_scrollable_frame.columnconfigure(1, weight=1)
        
        ui_labels_items = list(self.config.items("UI.Labels"))
        for r, (key, current_value) in enumerate(ui_labels_items):
            ttk.Label(ui_scrollable_frame, text=f"{key}:").grid(row=r, column=0, sticky=tk.W, padx=5, pady=3)
            entry = ttk.Entry(ui_scrollable_frame, width=60)
            entry.insert(0, current_value)
            entry.grid(row=r, column=1, sticky=tk.EW, padx=5, pady=3)
            self.ui_label_settings_entries[key] = entry
        ui_scrollable_frame.bind("<Configure>", lambda e: ui_canvas.itemconfig(ui_canvas_window, width=e.width))

        # Settings window buttons
        button_frame = ttk.Frame(settings_win, style='Custom.TFrame')
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=10)
        
        def save_settings_changes():
            logging.info("Saving settings changes.")
            
            # Save regular settings
            for (section, key), widget in self.settings_entries.items():
                if isinstance(widget, tk.BooleanVar):
                    self.config.set(section, key, str(widget.get()))
                elif hasattr(widget, 'get'):
                    self.config.set(section, key, widget.get())
            
            # Save UI labels
            for key, widget in self.ui_label_settings_entries.items():
                self.config.set("UI.Labels", key, widget.get())
            
            # Save external report folders
            folders = [self.ext_report_watch_folders_listbox.get(i) for i in range(self.ext_report_watch_folders_listbox.size())]
            self.config.set("Paths", "external_report_watch_folders_list", ';'.join(folders))
            
            # Save auto-send configurations for external reports
            for folder in folders:
                folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(folder)}"
                if not self.config.has_section(folder_section):
                    self.config.add_section(folder_section)
                
                # Only save if this folder is currently selected
                if (self._selected_ext_report_folder_for_autosend_config and 
                    self._selected_ext_report_folder_for_autosend_config == folder):
                    self.config.set(folder_section, "enabled", str(self.ext_report_autosend_widgets['enable_var'].get()))
                    self.config.set(folder_section, "recipients", self.ext_report_autosend_widgets['recipients'].get())
                    self.config.set(folder_section, "template", self.ext_report_autosend_widgets['template'].get())
                    self.config.set(folder_section, "custom_subject", self.ext_report_autosend_widgets['custom_subject'].get())
                    self.config.set(folder_section, "custom_body", self.ext_report_autosend_widgets['custom_body'].get("1.0", tk.END).strip())
            
            # Save favorite recipients
            favorites = [self.favorite_recipients_listbox.get(i) for i in range(self.favorite_recipients_listbox.size())]
            self.config.set("EmailRecipients", "favorite_list", ';'.join(favorites))
            
            save_config(self.config)
            
            # Apply changes
            self.apply_theme_and_styles()
            self.apply_ui_size()
            self.load_combobox_values_from_db()
            self._start_realtime_watchers()
            
            # Restart MWL server if configuration changed
            if PYNETDICOM_AVAILABLE:
                mwl_enabled = self.app_mwl_config_widgets['enabled_var'].get()
                current_mwl_enabled = self.config.getboolean("MWLServerConfig", "enabled", fallback=False)
                if mwl_enabled != current_mwl_enabled:
                    if mwl_enabled:
                        self.start_mwl_server_if_configured()
                    else:
                        self.stop_mwl_server()
            
            messagebox.showinfo("Settings Saved", "Settings have been saved and applied.", parent=settings_win)
            settings_win.destroy()
        
        ttk.Button(button_frame, text="Save Settings", command=save_settings_changes).pack(side=tk.RIGHT, padx=(5,0))
        ttk.Button(button_frame, text="Cancel", command=settings_win.destroy).pack(side=tk.RIGHT)

        # Populate auto-send template combobox
        template_names = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
        template_names = [t.strip() for t in template_names if t.strip()]
        if template_names:
            template_names.append("Custom")
            self.ext_report_autosend_widgets['template']['values'] = template_names
        
        self._set_autosend_widgets_state("disabled")

    def _setup_app_mwl_server_tab(self, tab):
        ttk.Label(tab, text=self.get_ui_label("app_mwl_server_tab_title", "MWL Server Configuration:"), 
                 font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(5,10), anchor=tk.W)
        
        if not PYNETDICOM_AVAILABLE:
            warning_frame = ttk.Frame(tab, style='Custom.TFrame')
            warning_frame.pack(fill=tk.X, pady=(0,10))
            warning_label = ttk.Label(warning_frame, 
                                     text="⚠️ MWL Server functionality is disabled because pynetdicom/pydicom libraries are not available.",
                                     foreground="red", font=('Helvetica', 11, 'bold'))
            warning_label.pack(anchor=tk.W)

        mwl_frame = ttk.Frame(tab, style='Custom.TFrame')
        mwl_frame.pack(expand=True, fill=tk.BOTH)
        mwl_frame.columnconfigure(1, weight=1)
        
        # Enable checkbox
        self.app_mwl_config_widgets['enabled_var'] = tk.BooleanVar(value=self.config.getboolean("MWLServerConfig", "enabled", fallback=False))
        enabled_check = ttk.Checkbutton(mwl_frame, text=self.get_ui_label("app_mwl_server_enable_label", "Enable MWL Server"),
                                       variable=self.app_mwl_config_widgets['enabled_var'],
                                       state="normal" if PYNETDICOM_AVAILABLE else "disabled")
        enabled_check.grid(row=0, column=0, columnspan=2, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("MWLServerConfig", "enabled")] = self.app_mwl_config_widgets['enabled_var']
        
        # AE Title
        ttk.Label(mwl_frame, text=self.get_ui_label("app_mwl_server_ae_title_label", "AE Title:")).grid(row=1, column=0, sticky=tk.W, padx=5, pady=7)
        ae_entry = ttk.Entry(mwl_frame, width=20, state="normal" if PYNETDICOM_AVAILABLE else "disabled")
        ae_entry.insert(0, self.config.get("MWLServerConfig", "ae_title", fallback="PYREGMWL"))
        ae_entry.grid(row=1, column=1, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("MWLServerConfig", "ae_title")] = ae_entry
        
        # Port
        ttk.Label(mwl_frame, text=self.get_ui_label("app_mwl_server_port_label", "Port:")).grid(row=2, column=0, sticky=tk.W, padx=5, pady=7)
        port_entry = ttk.Entry(mwl_frame, width=10, state="normal" if PYNETDICOM_AVAILABLE else "disabled")
        port_entry.insert(0, self.config.get("MWLServerConfig", "port", fallback="11112"))
        port_entry.grid(row=2, column=1, sticky=tk.W, padx=5, pady=7)
        self.settings_entries[("MWLServerConfig", "port")] = port_entry

    def _set_autosend_widgets_state(self, state):
        """Enable or disable auto-send configuration widgets"""
        for widget_key in ['enable', 'recipients', 'template']:
            if widget_key in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets[widget_key].config(state=state)
        
        self._toggle_autosend_custom_fields_state()

    def _toggle_autosend_custom_fields_state(self, event=None):
        """Show/hide custom subject and body fields based on template selection"""
        selected_template = self.ext_report_autosend_widgets.get('template', {}).get() if 'template' in self.ext_report_autosend_widgets else ""
        
        if selected_template == "Custom":
            # Show custom fields
            if 'custom_subject_label' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_subject_label'].pack(anchor=tk.W, pady=(5,0))
            if 'custom_subject' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_subject'].pack(fill=tk.X, pady=(0,5))
            if 'custom_body_label' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_body_label'].pack(anchor=tk.W, pady=(5,0))
            if 'custom_body' in self.ext_report_autosend_widgets:
                self.ext_report_autosend_widgets['custom_body'].pack(fill=tk.BOTH, expand=True, pady=(0,5))
        else:
            # Hide custom fields
            for widget_key in ['custom_subject_label', 'custom_subject', 'custom_body_label', 'custom_body']:
                if widget_key in self.ext_report_autosend_widgets:
                    self.ext_report_autosend_widgets[widget_key].pack_forget()

    def _add_favorite_to_autosend_recipients(self, parent_win):
        """Add a favorite recipient to the auto-send recipients field"""
        favorites_str = self.config.get("EmailRecipients", "favorite_list", fallback="")
        favorites = [e.strip() for e in favorites_str.split(';') if e.strip()]
        
        if not favorites:
            messagebox.showinfo("No Favorites", "No favorite recipients configured.", parent=parent_win)
            return
        
        def on_add():
            selected = favorite_listbox.curselection()
            if selected:
                email = favorite_listbox.get(selected[0])
                current_recipients = self.ext_report_autosend_widgets['recipients'].get()
                if current_recipients:
                    new_recipients = f"{current_recipients}; {email}"
                else:
                    new_recipients = email
                self.ext_report_autosend_widgets['recipients'].delete(0, tk.END)
                self.ext_report_autosend_widgets['recipients'].insert(0, new_recipients)
                add_win.destroy()
        
        add_win = tk.Toplevel(parent_win)
        add_win.title("Select Favorite Recipient")
        add_win.geometry("400x300")
        add_win.transient(parent_win)
        add_win.grab_set()
        
        ttk.Label(add_win, text="Select a favorite recipient to add:").pack(pady=10)
        
        favorite_listbox = tk.Listbox(add_win, selectmode=tk.SINGLE)
        favorite_listbox.pack(expand=True, fill=tk.BOTH, padx=10, pady=(0,10))
        
        for email in favorites:
            favorite_listbox.insert(tk.END, email)
        
        button_frame = ttk.Frame(add_win)
        button_frame.pack(pady=10)
        ttk.Button(button_frame, text="Add", command=on_add).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=add_win.destroy).pack(side=tk.LEFT, padx=5)

    def _on_ext_report_folder_select(self, event=None):
        """FIXED: Load auto-send configuration when a folder is selected"""
        selection = self.ext_report_watch_folders_listbox.curselection()
        if not selection:
            self._set_autosend_widgets_state("disabled")
            self._selected_ext_report_folder_for_autosend_config = None
            return
        
        selected_folder = self.ext_report_watch_folders_listbox.get(selection[0])
        self._selected_ext_report_folder_for_autosend_config = selected_folder
        self._set_autosend_widgets_state("normal")
        
        # Load configuration for this folder
        folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(selected_folder)}"
        
        # Populate template combobox first
        template_names = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
        template_names = [t.strip() for t in template_names if t.strip()]
        template_names.append("Custom")
        self.ext_report_autosend_widgets['template']['values'] = template_names
        
        if self.config.has_section(folder_section):
            # Load saved configuration
            enabled = self.config.getboolean(folder_section, "enabled", fallback=False)
            recipients = self.config.get(folder_section, "recipients", fallback="")
            template = self.config.get(folder_section, "template", fallback="")
            custom_subject = self.config.get(folder_section, "custom_subject", fallback="")
            custom_body = self.config.get(folder_section, "custom_body", fallback="")
            
            self.ext_report_autosend_widgets['enable_var'].set(enabled)
            self.ext_report_autosend_widgets['recipients'].delete(0, tk.END)
            self.ext_report_autosend_widgets['recipients'].insert(0, recipients)
            
            if template in template_names:
                self.ext_report_autosend_widgets['template'].set(template)
            else:
                self.ext_report_autosend_widgets['template'].set("")
            
            self.ext_report_autosend_widgets['custom_subject'].delete(0, tk.END)
            self.ext_report_autosend_widgets['custom_subject'].insert(0, custom_subject)
            
            self.ext_report_autosend_widgets['custom_body'].delete("1.0", tk.END)
            self.ext_report_autosend_widgets['custom_body'].insert("1.0", custom_body)
        else:
            # Reset to defaults
            self.ext_report_autosend_widgets['enable_var'].set(False)
            self.ext_report_autosend_widgets['recipients'].delete(0, tk.END)
            self.ext_report_autosend_widgets['template'].set("")
            self.ext_report_autosend_widgets['custom_subject'].delete(0, tk.END)
            self.ext_report_autosend_widgets['custom_body'].delete("1.0", tk.END)
        
        self._toggle_autosend_custom_fields_state()

    def _add_ext_report_folder(self, parent_win):
        folder = filedialog.askdirectory(title="Select External Report Watch Folder", parent=parent_win)
        if folder:
            folder = os.path.normpath(folder)
            # Check if already exists
            existing_folders = [self.ext_report_watch_folders_listbox.get(i) for i in range(self.ext_report_watch_folders_listbox.size())]
            if folder not in existing_folders:
                self.ext_report_watch_folders_listbox.insert(tk.END, folder)
                logging.info(f"Added external report watch folder: {folder}")
            else:
                messagebox.showinfo("Duplicate Folder", "This folder is already in the list.", parent=parent_win)

    def _remove_ext_report_folder(self):
        selection = self.ext_report_watch_folders_listbox.curselection()
        if selection:
            folder = self.ext_report_watch_folders_listbox.get(selection[0])
            self.ext_report_watch_folders_listbox.delete(selection[0])
            
            # Remove configuration section for this folder
            folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(folder)}"
            if self.config.has_section(folder_section):
                self.config.remove_section(folder_section)
            
            self._set_autosend_widgets_state("disabled")
            self._selected_ext_report_folder_for_autosend_config = None
            logging.info(f"Removed external report watch folder: {folder}")

    def _add_favorite_recipient(self, parent_win):
        email = simpledialog.askstring("Add Favorite Recipient", "Enter email address:", parent=parent_win)
        if email and email.strip():
            email = email.strip()
            # Basic email validation
            if "@" in email and "." in email.split("@")[-1]:
                # Check if already exists
                existing = [self.favorite_recipients_listbox.get(i) for i in range(self.favorite_recipients_listbox.size())]
                if email not in existing:
                    self.favorite_recipients_listbox.insert(tk.END, email)
                else:
                    messagebox.showinfo("Duplicate Email", "This email is already in the favorites list.", parent=parent_win)
            else:
                messagebox.showerror("Invalid Email", "Please enter a valid email address.", parent=parent_win)

    def _remove_favorite_recipient(self):
        selection = self.favorite_recipients_listbox.curselection()
        if selection:
            self.favorite_recipients_listbox.delete(selection[0])

    def _populate_email_templates_listbox(self):
        """Populate the email templates listbox"""
        self.email_templates_listbox.delete(0, tk.END)
        template_names = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
        for template in [t.strip() for t in template_names if t.strip()]:
            self.email_templates_listbox.insert(tk.END, template)

    def _clear_template_editor_fields(self):
        """Clear the template editor fields"""
        for widget in self.current_template_widgets.values():
            if isinstance(widget, ttk.Entry):
                widget.config(state="normal")
                widget.delete(0, tk.END)
                widget.config(state="readonly")
            elif isinstance(widget, tk.Text):
                widget.config(state="normal")
                widget.delete("1.0", tk.END)
                widget.config(state="disabled")

    def _load_template_for_editing(self, template_name_to_load=None):
        """Load selected template into the editor fields"""
        if template_name_to_load is None:
            selection = self.email_templates_listbox.curselection()
            if not selection:
                self._clear_template_editor_fields()
                return
            template_name_to_load = self.email_templates_listbox.get(selection[0])
        
        self._selected_template_for_edit = template_name_to_load
        
        # Load template data
        subject = self.config.get("EmailTemplates", f"{template_name_to_load}_subject", fallback="")
        body = self.config.get("EmailTemplates", f"{template_name_to_load}_body", fallback="")
        
        # Populate fields
        self.current_template_widgets['name'].config(state="normal")
        self.current_template_widgets['name'].delete(0, tk.END)
        self.current_template_widgets['name'].insert(0, template_name_to_load)
        self.current_template_widgets['name'].config(state="readonly")
        
        self.current_template_widgets['subject'].config(state="normal")
        self.current_template_widgets['subject'].delete(0, tk.END)
        self.current_template_widgets['subject'].insert(0, subject)
        
        self.current_template_widgets['body'].config(state="normal")
        self.current_template_widgets['body'].delete("1.0", tk.END)
        self.current_template_widgets['body'].insert("1.0", body)

    def _edit_email_template(self, parent_win, is_new=False):
        """Open dialog to edit or create email template"""
        if is_new:
            template_name = simpledialog.askstring("New Template", "Enter template name:", parent=parent_win)
            if not template_name or not template_name.strip():
                return
            template_name = template_name.strip()
            
            # Check if already exists
            existing_templates = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
            existing_templates = [t.strip() for t in existing_templates if t.strip()]
            if template_name in existing_templates:
                messagebox.showerror("Duplicate Template", "A template with this name already exists.", parent=parent_win)
                return
            
            # Add to templates list
            if existing_templates:
                new_template_list = ';'.join(existing_templates + [template_name])
            else:
                new_template_list = template_name
            self.config.set("EmailTemplates", "template_names", new_template_list)
            
            # Set default values
            self.config.set("EmailTemplates", f"{template_name}_subject", "Report: {Patient Name}")
            self.config.set("EmailTemplates", f"{template_name}_body", "Please find attached the report for {Patient Name}.")
            
            # Refresh listbox and select new template
            self._populate_email_templates_listbox()
            for i in range(self.email_templates_listbox.size()):
                if self.email_templates_listbox.get(i) == template_name:
                    self.email_templates_listbox.selection_set(i)
                    break
            self._load_template_for_editing(template_name)
        
        # Enable editing
        if hasattr(self, '_selected_template_for_edit') and self._selected_template_for_edit:
            self.current_template_widgets['subject'].config(state="normal")
            self.current_template_widgets['body'].config(state="normal")

    def _save_current_template_changes(self):
        """Save changes to the currently selected template"""
        if not hasattr(self, '_selected_template_for_edit') or not self._selected_template_for_edit:
            messagebox.showwarning("No Template Selected", "Please select a template to edit.")
            return
        
        template_name = self._selected_template_for_edit
        subject = self.current_template_widgets['subject'].get()
        body = self.current_template_widgets['body'].get("1.0", tk.END).strip()
        
        self.config.set("EmailTemplates", f"{template_name}_subject", subject)
        self.config.set("EmailTemplates", f"{template_name}_body", body)
        
        # Disable editing
        self.current_template_widgets['subject'].config(state="readonly")
        self.current_template_widgets['body'].config(state="disabled")
        
        messagebox.showinfo("Template Saved", f"Template '{template_name}' has been saved.")

    def _delete_email_template(self):
        """Delete the selected email template"""
        selection = self.email_templates_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a template to delete.")
            return
        
        template_name = self.email_templates_listbox.get(selection[0])
        
        if template_name == "Default":
            messagebox.showerror("Cannot Delete", "The Default template cannot be deleted.")
            return
        
        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete the template '{template_name}'?"):
            # Remove from templates list
            existing_templates = self.config.get("EmailTemplates", "template_names", fallback="").split(';')
            existing_templates = [t.strip() for t in existing_templates if t.strip() and t != template_name]
            self.config.set("EmailTemplates", "template_names", ';'.join(existing_templates))
            
            # Remove template data
            self.config.remove_option("EmailTemplates", f"{template_name}_subject")
            self.config.remove_option("EmailTemplates", f"{template_name}_body")
            
            # Refresh listbox and clear editor
            self._populate_email_templates_listbox()
            self._clear_template_editor_fields()
            self._selected_template_for_edit = None

    def browse_path(self, entry_widget, is_folder, parent=None):
        if is_folder:
            path = filedialog.askdirectory(title="Select Folder", parent=parent or self.root)
        else:
            path = filedialog.askopenfilename(title="Select File", parent=parent or self.root)
        
        if path:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, os.path.normpath(path))

    def view_patient_data_window(self):
        logging.info("Opening patient data viewer window.")
        data_win = tk.Toplevel(self.root)
        data_win.title(self.get_ui_label("view_data_window_title", "Patient Data Viewer"))
        data_win.geometry("1200x700")
        data_win.transient(self.root)
        data_win.configure(bg=self.current_palette.get("bg", "#F0F0F0"))

        search_frame = ttk.Frame(data_win, style='Custom.TFrame', padding=10)
        search_frame.pack(side=tk.TOP, fill=tk.X)
        ttk.Label(search_frame, text="Search:").pack(side=tk.LEFT, padx=(0,5))
        search_entry = ttk.Entry(search_frame, width=30)
        search_entry.pack(side=tk.LEFT, padx=(0,10))
        
        tree_frame = ttk.Frame(data_win)
        tree_frame.pack(expand=True, fill=tk.BOTH, padx=10, pady=(0,10))
        
        columns = ("id", "patient_name", "patient_id", "accession_number", "dob_yyyymmdd", "sex", 
                  "study_date", "study_time", "study_description", "referred_from", "modality", 
                  "requesting_physician", "requested_procedure_id", "scheduled_station_ae_title", "created_at")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=20)
        
        column_widths = {"id": 50, "patient_name": 150, "patient_id": 100, "accession_number": 120, "dob_yyyymmdd": 90, 
                        "sex": 40, "study_date": 80, "study_time": 70, "study_description": 200, "referred_from": 120, 
                        "modality": 70, "requesting_physician": 150, "requested_procedure_id": 120, "scheduled_station_ae_title": 100, "created_at": 130}
        
        for col in columns:
            tree.heading(col, text=col.replace("_", " ").title(), command=lambda c=col: self.sort_treeview_column(tree, c, False))
            tree.column(col, width=column_widths.get(col, 100), minwidth=50)
        
        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=tree_scroll.set)
        tree.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        def filter_treeview_data(event=None):
            search_term = search_entry.get().strip()
            tree.delete(*tree.get_children())
            
            try:
                all_data, _ = get_all_patient_records_db(search_term)
                for record in all_data:
                    values = [record.get(col, "") for col in columns]
                    tree.insert("", tk.END, values=values)
                logging.info(f"Loaded {len(all_data)} patient records into viewer (search: '{search_term}').")
            except Exception as e:
                messagebox.showerror("Database Error", f"Error loading patient data: {e}", parent=data_win)
                logging.exception("Error loading patient data for viewer.")

        search_entry.bind("<KeyRelease>", filter_treeview_data)
        ttk.Button(search_frame, text="Refresh", command=filter_treeview_data).pack(side=tk.LEFT, padx=(10,0))
        
        filter_treeview_data()

    def sort_treeview_column(self, tv, col, reverse):
        data_list = [(tv.set(k, col), k) for k in tv.get_children('')]
        try:
            data_list.sort(key=lambda x: float(x[0]) if x[0].replace('.', '', 1).isdigit() else x[0].lower(), reverse=reverse)
        except (ValueError, AttributeError):
            data_list.sort(key=lambda x: str(x[0]).lower(), reverse=reverse)
        
        for index, (val, k) in enumerate(data_list):
            tv.move(k, '', index)
        
        tv.heading(col, command=lambda: self.sort_treeview_column(tv, col, not reverse))

    def update_email_button_state(self):
        try:
            all_data, _ = get_all_patient_records_db("")
            if all_data:
                self.email_button.config(state=tk.NORMAL)
            else:
                self.email_button.config(state=tk.DISABLED)
        except Exception as e:
            logging.exception("Error updating email button state.")
            self.email_button.config(state=tk.DISABLED)

    def trigger_email_report_picker(self):
        logging.info("Opening email report picker dialog.")
        ReportPickerDialog(self.root, self.config, self.current_palette, self.get_ui_label, self)

    def update_recent_recipients(self, recipient_email):
        current_recents_str = self.config.get("EmailRecipients", "recent_list", fallback="")
        current_recents = [e.strip() for e in current_recents_str.split(';') if e.strip()]
        
        if recipient_email in current_recents:
            current_recents.remove(recipient_email)
        current_recents.insert(0, recipient_email)
        
        max_recent = self.config.getint("EmailRecipients", "max_recent", fallback=10)
        current_recents = current_recents[:max_recent]
        
        self.config.set("EmailRecipients", "recent_list", ';'.join(current_recents))
        save_config(self.config)

    def send_email_with_report(self, recipient_emails_str, subject, body, attachment_paths, update_recents=True):
        smtp_server = self.config.get("SMTP", "server", fallback="")
        smtp_port = self.config.getint("SMTP", "port", fallback=587)
        smtp_user = self.config.get("SMTP", "user", fallback="")
        smtp_password = self.config.get("SMTP", "password", fallback="")
        sender_email = self.config.get("SMTP", "sender_email", fallback="")
        use_tls = self.config.getboolean("SMTP", "use_tls", fallback=True)

        if not smtp_server or not sender_email:
            messagebox.showerror("SMTP Configuration Error", "SMTP server and sender email must be configured in Settings.", parent=self.root)
            return False

        try:
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = recipient_emails_str
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))

            for attachment_path in attachment_paths:
                if os.path.isfile(attachment_path):
                    with open(attachment_path, "rb") as attachment:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(attachment.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', f'attachment; filename= {os.path.basename(attachment_path)}')
                        msg.attach(part)

            server = smtplib.SMTP(smtp_server, smtp_port)
            if use_tls:
                server.starttls()
            if smtp_user and smtp_password:
                server.login(smtp_user, smtp_password)
            
            text = msg.as_string()
            server.sendmail(sender_email, [e.strip() for e in recipient_emails_str.split(';')], text)
            server.quit()

            if update_recents:
                for recipient in [e.strip() for e in recipient_emails_str.split(';') if e.strip()]:
                    self.update_recent_recipients(recipient)
            
            self.update_status(f"Email sent successfully to: {recipient_emails_str}")
            logging.info(f"Email sent successfully to: {recipient_emails_str}")
            return True

        except Exception as e:
            error_msg = f"Error sending email: {e}"
            messagebox.showerror("Email Error", error_msg, parent=self.root)
            self.update_status(error_msg, True)
            logging.exception("Error sending email.")
            return False

    def test_smtp_settings(self, parent_window):
        smtp_server = self.settings_entries[("SMTP", "server")].get()
        smtp_port_str = self.settings_entries[("SMTP", "port")].get()
        smtp_user = self.settings_entries[("SMTP", "user")].get()
        smtp_password = self.settings_entries[("SMTP", "password")].get()
        sender_email = self.settings_entries[("SMTP", "sender_email")].get()
        use_tls = self.settings_entries[("SMTP", "use_tls")].get()

        if not smtp_server or not sender_email:
            messagebox.showerror("Configuration Error", "SMTP server and sender email are required for testing.", parent=parent_window)
            return

        try:
            smtp_port = int(smtp_port_str)
        except ValueError:
            messagebox.showerror("Configuration Error", "SMTP port must be a valid number.", parent=parent_window)
            return

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            if use_tls:
                server.starttls()
            if smtp_user and smtp_password:
                server.login(smtp_user, smtp_password)
            server.quit()
            
            messagebox.showinfo("SMTP Test Successful", "SMTP settings are working correctly!", parent=parent_window)
            logging.info("SMTP settings test successful.")
        except Exception as e:
            messagebox.showerror("SMTP Test Failed", f"SMTP test failed: {e}", parent=parent_window)
            logging.error(f"SMTP test failed: {e}")

    def get_all_pdf_watch_folders(self):
        folders = []
        for mod in MODALITIES:
            folder = self.get_modality_specific_path("Paths.WatchFolders.Modalities", mod)
            if folder and os.path.isdir(folder):
                folders.append(folder)
        
        general_folder = self.config.get("Paths", "general_watch_folder", fallback="")
        if general_folder and os.path.isdir(general_folder):
            folders.append(general_folder)
        
        return folders

    def get_all_external_report_watch_folders(self):
        folders_str = self.config.get("Paths", "external_report_watch_folders_list", fallback="")
        folders = []
        for folder in folders_str.split(';'):
            folder = folder.strip()
            if folder:
                normalized_folder = os.path.normpath(os.path.expanduser(folder))
                folders.append(normalized_folder)
        return folders

    def _process_auto_send_for_external_report_file(self, file_path):
        """FIXED: Process auto-send email for external report files with patient data extraction"""
        filename = os.path.basename(file_path)
        folder_path = os.path.dirname(file_path)
        
        logging.info(f"Processing auto-send for external report: {filename} in {folder_path}")
        
        # Find the folder configuration
        folder_section = f"ExternalReportAutoSend.{normalize_path_for_config_section(folder_path)}"
        
        if not self.config.has_section(folder_section):
            logging.info(f"No auto-send configuration found for folder: {folder_path}")
            return
        
        if not self.config.getboolean(folder_section, "enabled", fallback=False):
            logging.info(f"Auto-send disabled for folder: {folder_path}")
            return
        
        recipients = self.config.get(folder_section, "recipients", fallback="")
        if not recipients:
            logging.warning(f"No recipients configured for auto-send in folder: {folder_path}")
            return
        
        template_name = self.config.get(folder_section, "template", fallback="")
        
        # FIXED: Extract patient data from PDF filename
        patient_data = extract_patient_data_from_filename(filename)
        
        # Prepare email subject and body
        if template_name == "Custom":
            subject = self.config.get(folder_section, "custom_subject", fallback="External Report: {Filename}")
            body = self.config.get(folder_section, "custom_body", fallback="Please find attached the external report: {Filename}")
        elif template_name and self.config.has_option("EmailTemplates", f"{template_name}_subject"):
            subject = self.config.get("EmailTemplates", f"{template_name}_subject", fallback="External Report: {Filename}")
            body = self.config.get("EmailTemplates", f"{template_name}_body", fallback="Please find attached the external report: {Filename}")
        else:
            subject = "External Report: {Filename}"
            body = "Please find attached the external report: {Filename}"
        
        # FIXED: Replace placeholders with extracted patient data and file info
        current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        replacements = {
            '{Filename}': filename,
            '{FolderPath}': folder_path,
            '{DateTime}': current_datetime,
            '{Patient Name}': patient_data.get('patient_name', 'Unknown Patient'),
            '{Modality}': patient_data.get('modality', 'Unknown'),
            '{Study Description}': patient_data.get('description', 'Report'),
            '{Date}': patient_data.get('date', datetime.now().strftime("%Y-%m-%d")),
            '{Report Filename}': filename,
            '{Attachment Count}': '1',
            '{All Attachment Names}': filename
        }
        
        for placeholder, value in replacements.items():
            subject = subject.replace(placeholder, str(value))
            body = body.replace(placeholder, str(value))
        
        # Send the email
        try:
            success = self.send_email_with_report(recipients, subject, body, [file_path], update_recents=False)
            if success:
                self.update_status(f"Auto-sent external report: {filename}")
                logging.info(f"Successfully auto-sent external report: {filename} to {recipients}")
            else:
                logging.error(f"Failed to auto-send external report: {filename}")
        except Exception as e:
            logging.exception(f"Error in auto-send for external report {filename}: {e}")

    def open_served_worklist_viewer(self):
        """FIXED: Open worklist viewer with save functionality"""
        logging.info("Opening served worklist viewer.")
        ServedWorklistDialog(self.root, self.config, self.current_palette, self.get_ui_label, self)


# --- ReportPickerDialog Class ---
class ReportPickerDialog(tk.Toplevel):
    def __init__(self, parent, app_config, palette, get_ui_label_func, main_app_ref):
        super().__init__(parent)
        self.app_config = app_config
        self.palette = palette
        self.get_ui_label = get_ui_label_func
        self.main_app = main_app_ref
        
        self.title(self.get_ui_label("email_picker_title", "Select Report to Email"))
        self.geometry("900x600")
        self.transient(parent)
        self.grab_set()
        self.configure(bg=self.palette.get("bg", "#F0F0F0"))
        
        main_frame = ttk.Frame(self, style='Custom.TFrame', padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        ttk.Label(main_frame, text="Select a report to email:", font=('Helvetica', 12, 'bold'), style="Header.TLabel").pack(pady=(0,10), anchor=tk.W)
        
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(expand=True, fill=tk.BOTH, pady=(0,15))
        
        columns = ("patient_name", "patient_id", "accession_number", "study_description", "modality", "study_date", "docx_path")
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        
        column_widths = {"patient_name": 150, "patient_id": 100, "accession_number": 120, "study_description": 200, "modality": 80, "study_date": 100, "docx_path": 250}
        for col in columns:
            self.tree.heading(col, text=col.replace("_", " ").title(), command=lambda c=col: self._sort_column(self.tree, c, False))
            self.tree.column(col, width=column_widths.get(col, 100), minwidth=50)
        
        tree_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scroll.set)
        self.tree.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tree.bind("<<TreeviewSelect>>", self.on_tree_select)
        self.tree.bind("<Double-1>", self.on_compose_email)
        
        button_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10,0))
        
        self.compose_button = ttk.Button(button_frame, text=self.get_ui_label("email_picker_button", "Compose Email for Selected Report"), 
                                        command=self.on_compose_email, state=tk.DISABLED)
        self.compose_button.pack(side=tk.LEFT, padx=(0,10))
        
        ttk.Button(button_frame, text="Cancel", command=self.on_cancel).pack(side=tk.RIGHT)
        ttk.Button(button_frame, text="Refresh", command=self.load_reports).pack(side=tk.RIGHT, padx=(0,10))
        
        self.load_reports()

    def _sort_column(self, tv, col, reverse):
        data_list = [(tv.set(k, col), k) for k in tv.get_children('')]
        try:
            data_list.sort(key=lambda x: float(x[0]) if x[0].replace('.', '', 1).isdigit() else x[0].lower(), reverse=reverse)
        except (ValueError, AttributeError):
            data_list.sort(key=lambda x: str(x[0]).lower(), reverse=reverse)
        
        for index, (val, k) in enumerate(data_list):
            tv.move(k, '', index)
        
        tv.heading(col, command=lambda: self._sort_column(tv, col, not reverse))

    def load_reports(self):
        self.tree.delete(*self.tree.get_children())
        
        try:
            all_data, _ = get_all_patient_records_db("")
            report_count = 0
            
            for record in all_data:
                patient_name = record.get("patient_name", "Unknown")
                patient_id = record.get("patient_id", "Unknown") 
                accession_number = record.get("accession_number", "Unknown")
                study_description = record.get("study_description", "Unknown")
                modality = record.get("modality", "Unknown")
                study_date = format_date_friendly(record.get("study_date", ""))
                
                modality_output_folder = self.main_app.get_modality_specific_path("Paths.Output.DOCX.Modalities", modality)
                if not modality_output_folder:
                    continue
                
                safe_pname = "".join(c if c.isalnum() else "_" for c in patient_name)
                safe_pid = patient_id.replace(' ', '_')
                patient_subfolder_name = f"{safe_pname}_{safe_pid}"
                patient_folder = os.path.join(modality_output_folder, patient_subfolder_name)
                
                if os.path.isdir(patient_folder):
                    for filename in os.listdir(patient_folder):
                        if filename.endswith("_REPORT.docx"):
                            docx_path = os.path.join(patient_folder, filename)
                            values = (patient_name, patient_id, accession_number, study_description, modality, study_date, docx_path)
                            self.tree.insert("", tk.END, values=values)
                            report_count += 1
            
            logging.info(f"Loaded {report_count} DOCX reports into email picker.")
            
        except Exception as e:
            messagebox.showerror("Error Loading Reports", f"Error loading reports: {e}", parent=self)
            logging.exception("Error loading reports for email picker.")

    def on_tree_select(self, event=None):
        self.compose_button.config(state=tk.NORMAL if self.tree.selection() else tk.DISABLED)

    def on_compose_email(self, event=None):
        selection = self.tree.selection()
        if not selection:
            return
        
        item = self.tree.item(selection[0])
        values = item['values']
        
        if len(values) < 7:
            messagebox.showerror("Selection Error", "Invalid report selection.", parent=self)
            return
        
        patient_name, patient_id, accession_number, study_description, modality, study_date, docx_path = values
        
        if not os.path.exists(docx_path):
            messagebox.showerror("File Not Found", f"Report file not found:\n{docx_path}", parent=self)
            return
        
        report_folder = os.path.dirname(docx_path)
        attachment_files = [docx_path]
        
        for filename in os.listdir(report_folder):
            file_path = os.path.join(report_folder, filename)
            if os.path.isfile(file_path) and file_path != docx_path:
                attachment_files.append(file_path)
        
        self.destroy()
        EmailComposerDialog(self.master, self.app_config, self.palette, self.get_ui_label, self.main_app,
                          patient_name=patient_name, modality=modality, study_description=study_description,
                          study_date=study_date, attachment_files=attachment_files)

    def on_cancel(self):
        self.destroy()


# --- EmailComposerDialog Class ---
class EmailComposerDialog(tk.Toplevel):
    def __init__(self, parent, app_config, palette, get_ui_label_func, main_app_ref, 
                 patient_name="", modality="", study_description="", study_date="", attachment_files=None):
        super().__init__(parent)
        self.app_config = app_config
        self.palette = palette
        self.get_ui_label = get_ui_label_func
        self.main_app = main_app_ref
        self.patient_name = patient_name
        self.modality = modality
        self.study_description = study_description
        self.study_date = study_date
        self.attachment_files = attachment_files or []
        
        self.title(self.get_ui_label("email_composer_title", "Compose Email"))
        self.geometry("800x700")
        self.transient(parent)
        self.grab_set()
        self.configure(bg=self.palette.get("bg", "#F0F0F0"))
        
        main_frame = ttk.Frame(self, style='Custom.TFrame', padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)
        main_frame.columnconfigure(1, weight=1)
        
        # Recipients section
        recipients_frame = ttk.LabelFrame(main_frame, text=self.get_ui_label("email_composer_select_recipients_label", "Select Recipients:"), 
                                        style='Custom.TFrame', padding=10)
        recipients_frame.grid(row=0, column=0, columnspan=2, sticky=tk.EW, pady=(0,10))
        recipients_frame.columnconfigure(1, weight=1)
        
        ttk.Label(recipients_frame, text=self.get_ui_label("recent_recipients_combobox_label", "Recent/Favorites:")).grid(row=0, column=0, sticky=tk.W, padx=(0,5))
        self.recipients_listbox = tk.Listbox(recipients_frame, selectmode=tk.MULTIPLE, height=4,
                                           bg=self.palette.get("entry_bg"), fg=self.palette.get("entry_fg"),
                                           selectbackground=self.palette.get("header_fg"), selectforeground=self.palette.get("button_fg"))
        self.recipients_listbox.grid(row=0, column=1, sticky=tk.EW, padx=5)
        
        ttk.Button(recipients_frame, text=self.get_ui_label("email_composer_add_selected_button", "Add Selected to 'To:' Field"), 
                  command=self._add_selected_recipients_to_to_field).grid(row=0, column=2, padx=(5,0))
        
        self.populate_recipient_listbox()
        
        # To field
        ttk.Label(main_frame, text=self.get_ui_label("email_composer_to", "To:")).grid(row=1, column=0, sticky=tk.W, pady=(5,0))
        self.to_entry = ttk.Entry(main_frame, width=60)
        self.to_entry.grid(row=1, column=1, sticky=tk.EW, padx=(10,0), pady=(5,0))
        
        # Template selection
        template_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        template_frame.grid(row=2, column=0, columnspan=2, sticky=tk.EW, pady=(10,0))
        template_frame.columnconfigure(1, weight=1)
        
        ttk.Label(template_frame, text=self.get_ui_label("email_composer_template_label", "Email Template:")).pack(side=tk.LEFT)
        self.template_combo = ttk.Combobox(template_frame, width=30, state="readonly")
        self.template_combo.pack(side=tk.LEFT, padx=(10,0))
        self.template_combo.bind("<<ComboboxSelected>>", self.on_template_selected)
        
        self.populate_template_combobox()
        
        # Subject field
        ttk.Label(main_frame, text=self.get_ui_label("email_composer_subject", "Subject:")).grid(row=3, column=0, sticky=tk.W, pady=(10,0))
        self.subject_entry = ttk.Entry(main_frame, width=60)
        self.subject_entry.grid(row=3, column=1, sticky=tk.EW, padx=(10,0), pady=(10,0))
        
        # Body field
        ttk.Label(main_frame, text=self.get_ui_label("email_composer_body", "Body:")).grid(row=4, column=0, sticky=tk.NW, pady=(10,0))
        self.body_text = tk.Text(main_frame, height=15, width=60, 
                                bg=self.palette.get("entry_bg"), fg=self.palette.get("entry_fg"),
                                insertbackground=self.palette.get("entry_fg"))
        self.body_text.grid(row=4, column=1, sticky=tk.EW + tk.NS, padx=(10,0), pady=(10,0))
        main_frame.rowconfigure(4, weight=1)
        
        # Attachments info
        attachments_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        attachments_frame.grid(row=5, column=0, columnspan=2, sticky=tk.EW, pady=(10,0))
        
        ttk.Label(attachments_frame, text=self.get_ui_label("email_composer_attachments_label", "Attachments:")).pack(side=tk.LEFT)
        attachment_names = [os.path.basename(f) for f in self.attachment_files]
        ttk.Label(attachments_frame, text=f"{len(attachment_names)} file(s): {', '.join(attachment_names[:3])}{'...' if len(attachment_names) > 3 else ''}",
                 foreground="gray").pack(side=tk.LEFT, padx=(10,0))
        
        # Buttons
        button_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        button_frame.grid(row=6, column=0, columnspan=2, sticky=tk.EW, pady=(15,0))
        
        ttk.Button(button_frame, text=self.get_ui_label("email_composer_send_button", "Send Email"), 
                  command=self.send_composed_email).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Cancel", command=self.destroy).pack(side=tk.RIGHT)
        
        # Apply default template if available
        if self.template_combo['values']:
            self.template_combo.set(self.template_combo['values'][0])
            self.apply_email_template()

    def populate_recipient_listbox(self):
        recent_str = self.app_config.get("EmailRecipients", "recent_list", fallback="")
        favorite_str = self.app_config.get("EmailRecipients", "favorite_list", fallback="")
        
        recent_emails = [e.strip() for e in recent_str.split(';') if e.strip()]
        favorite_emails = [e.strip() for e in favorite_str.split(';') if e.strip()]
        
        all_emails = []
        for email in favorite_emails:
            all_emails.append(f"⭐ {email}")
        for email in recent_emails:
            if email not in favorite_emails:
                all_emails.append(f"📧 {email}")
        
        for email in all_emails:
            self.recipients_listbox.insert(tk.END, email)

    def _add_selected_recipients_to_to_field(self):
        selection = self.recipients_listbox.curselection()
        selected_emails = []
        
        for idx in selection:
            item = self.recipients_listbox.get(idx)
            # Remove the emoji prefix
            email = item.replace("⭐ ", "").replace("📧 ", "")
            selected_emails.append(email)
        
        if selected_emails:
            current_to = self.to_entry.get().strip()
            if current_to:
                new_to = f"{current_to}; {'; '.join(selected_emails)}"
            else:
                new_to = '; '.join(selected_emails)
            
            self.to_entry.delete(0, tk.END)
            self.to_entry.insert(0, new_to)

    def populate_template_combobox(self):
        template_names = self.app_config.get("EmailTemplates", "template_names", fallback="").split(';')
        template_names = [t.strip() for t in template_names if t.strip()]
        if template_names:
            self.template_combo['values'] = template_names

    def on_template_selected(self, event=None):
        self.apply_email_template()

    def apply_email_template(self):
        template_name = self.template_combo.get()
        if not template_name:
            return
        
        subject_template = self.app_config.get("EmailTemplates", f"{template_name}_subject", fallback="")
        body_template = self.app_config.get("EmailTemplates", f"{template_name}_body", fallback="")
        
        # Replace placeholders
        replacements = {
            '{Patient Name}': self.patient_name,
            '{Modality}': self.modality,
            '{Study Description}': self.study_description,
            '{Date}': self.study_date,
            '{Report Filename}': os.path.basename(self.attachment_files[0]) if self.attachment_files else "",
            '{Attachment Count}': str(len(self.attachment_files)),
            '{All Attachment Names}': ', '.join([os.path.basename(f) for f in self.attachment_files])
        }
        
        for placeholder, value in replacements.items():
            subject_template = subject_template.replace(placeholder, str(value))
            body_template = body_template.replace(placeholder, str(value))
        
        self.subject_entry.delete(0, tk.END)
        self.subject_entry.insert(0, subject_template)
        
        self.body_text.delete("1.0", tk.END)
        self.body_text.insert("1.0", body_template)

    def send_composed_email(self):
        to_emails = self.to_entry.get().strip()
        subject = self.subject_entry.get().strip()
        body = self.body_text.get("1.0", tk.END).strip()
        
        if not to_emails:
            messagebox.showerror("Missing Recipients", "Please enter at least one recipient email address.", parent=self)
            return
        
        if not subject:
            messagebox.showerror("Missing Subject", "Please enter an email subject.", parent=self)
            return
        
        success = self.main_app.send_email_with_report(to_emails, subject, body, self.attachment_files)
        if success:
            messagebox.showinfo("Email Sent", "Email sent successfully!", parent=self)
            self.destroy()


# --- ServedWorklistDialog Class ---
class ServedWorklistDialog(tk.Toplevel):
    def __init__(self, parent, app_config, palette, get_ui_label_func, main_app_ref):
        super().__init__(parent)
        self.app_config = app_config
        self.palette = palette
        self.get_ui_label = get_ui_label_func
        self.main_app = main_app_ref
        self.original_data = {}  # Store original data for change detection
        self.changes_made = False
        
        self.title(self.get_ui_label("view_served_worklist_title", "Served Worklist Viewer"))
        self.geometry("1400x800")
        self.transient(parent)
        self.configure(bg=self.palette.get("bg", "#F0F0F0"))
        
        main_frame = ttk.Frame(self, style='Custom.TFrame', padding=15)
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        # Header
        header_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        header_frame.pack(fill=tk.X, pady=(0,10))
        
        ttk.Label(header_frame, text="MWL Server Database Entries", 
                 font=('Helvetica', 14, 'bold'), style="Header.TLabel").pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text="Refresh", command=self.refresh_list).pack(side=tk.RIGHT, padx=(10,0))
        
        # FIXED: Add save button (initially hidden)
        self.save_button = ttk.Button(header_frame, text="Save Changes", command=self.save_changes, state=tk.DISABLED)
        self.save_button.pack(side=tk.RIGHT, padx=(10,0))
        
        # Tree view
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(expand=True, fill=tk.BOTH, pady=(0,10))
        
        columns = ("id", "patient_name", "patient_id", "accession_number", "dob_yyyymmdd", "sex", 
                  "study_date", "study_time", "study_description", "referred_from", "modality", 
                  "requesting_physician", "requested_procedure_id", "scheduled_station_ae_title")
        
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=20, selectmode="extended")
        
        column_widths = {"id": 50, "patient_name": 150, "patient_id": 100, "accession_number": 120, 
                        "dob_yyyymmdd": 90, "sex": 40, "study_date": 80, "study_time": 70, 
                        "study_description": 200, "referred_from": 120, "modality": 70, 
                        "requesting_physician": 150, "requested_procedure_id": 120, "scheduled_station_ae_title": 100}
        
        for col in columns:
            self.tree.heading(col, text=col.replace("_", " ").title())
            self.tree.column(col, width=column_widths.get(col, 100), minwidth=50)
        
        # FIXED: Make modality column read-only by setting it to a different color
        self.tree.tag_configure("readonly", background="#f0f0f0")
        
        tree_scroll_v = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        tree_scroll_h = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=tree_scroll_v.set, xscrollcommand=tree_scroll_h.set)
        
        self.tree.grid(row=0, column=0, sticky=tk.NSEW)
        tree_scroll_v.grid(row=0, column=1, sticky=tk.NS)
        tree_scroll_h.grid(row=1, column=0, sticky=tk.EW)
        
        tree_frame.rowconfigure(0, weight=1)
        tree_frame.columnconfigure(0, weight=1)
        
        # FIXED: Bind double-click to edit functionality
        self.tree.bind("<Double-1>", self.on_item_double_click)
        
        # Edit frame (initially hidden)
        self.edit_frame = ttk.LabelFrame(main_frame, text="Edit Record", style='Custom.TFrame', padding=10)
        self.edit_widgets = {}
        self.current_editing_id = None
        
        # Create edit widgets
        self.create_edit_widgets()
        
        # Buttons
        button_frame = ttk.Frame(main_frame, style='Custom.TFrame')
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10,0))
        
        # Add Delete button
        ttk.Button(button_frame, text="Delete Selected", command=self.delete_selected_records).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Close", command=self.destroy).pack(side=tk.RIGHT)
        
        self.refresh_list()
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def create_edit_widgets(self):
        """Create editing widgets for the selected record"""
        # Patient Name
        ttk.Label(self.edit_frame, text="Patient Name:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['patient_name'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['patient_name'].grid(row=0, column=1, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['patient_name'].bind("<KeyRelease>", self.on_field_change)
        
        # Patient ID
        ttk.Label(self.edit_frame, text="Patient ID:").grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['patient_id'] = ttk.Entry(self.edit_frame, width=20)
        self.edit_widgets['patient_id'].grid(row=0, column=3, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['patient_id'].bind("<KeyRelease>", self.on_field_change)
        
        # Accession Number
        ttk.Label(self.edit_frame, text="Accession Number:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['accession_number'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['accession_number'].grid(row=1, column=1, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['accession_number'].bind("<KeyRelease>", self.on_field_change)
        
        # Date of Birth
        ttk.Label(self.edit_frame, text="Date of Birth:").grid(row=1, column=2, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['dob_yyyymmdd'] = ttk.Entry(self.edit_frame, width=20)
        self.edit_widgets['dob_yyyymmdd'].grid(row=1, column=3, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['dob_yyyymmdd'].bind("<KeyRelease>", self.on_field_change)
        
        # Sex
        ttk.Label(self.edit_frame, text="Sex:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['sex'] = ttk.Combobox(self.edit_frame, width=10, values=["M", "F", "O"], state="readonly")
        self.edit_widgets['sex'].grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['sex'].bind("<<ComboboxSelected>>", self.on_field_change)
        
        # Study Date
        ttk.Label(self.edit_frame, text="Study Date:").grid(row=2, column=2, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['study_date'] = ttk.Entry(self.edit_frame, width=20)
        self.edit_widgets['study_date'].grid(row=2, column=3, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['study_date'].bind("<KeyRelease>", self.on_field_change)
        
        # Study Time
        ttk.Label(self.edit_frame, text="Study Time:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['study_time'] = ttk.Entry(self.edit_frame, width=20)
        self.edit_widgets['study_time'].grid(row=3, column=1, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['study_time'].bind("<KeyRelease>", self.on_field_change)
        
        # Study Description
        ttk.Label(self.edit_frame, text="Study Description:").grid(row=3, column=2, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['study_description'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['study_description'].grid(row=3, column=3, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['study_description'].bind("<KeyRelease>", self.on_field_change)
        
        # Referred From
        ttk.Label(self.edit_frame, text="Referred From:").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['referred_from'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['referred_from'].grid(row=4, column=1, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['referred_from'].bind("<KeyRelease>", self.on_field_change)
        
        # FIXED: Modality (read-only)
        ttk.Label(self.edit_frame, text="Modality:").grid(row=4, column=2, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['modality'] = ttk.Entry(self.edit_frame, width=20, state="readonly")
        self.edit_widgets['modality'].grid(row=4, column=3, sticky=tk.EW, padx=5, pady=5)
        
        # Requesting Physician
        ttk.Label(self.edit_frame, text="Requesting Physician:").grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['requesting_physician'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['requesting_physician'].grid(row=5, column=1, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['requesting_physician'].bind("<KeyRelease>", self.on_field_change)
        
        # Requested Procedure ID
        ttk.Label(self.edit_frame, text="Requested Procedure ID:").grid(row=5, column=2, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['requested_procedure_id'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['requested_procedure_id'].grid(row=5, column=3, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['requested_procedure_id'].bind("<KeyRelease>", self.on_field_change)
        
        # Scheduled Station AE Title
        ttk.Label(self.edit_frame, text="Scheduled Station AE:").grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
        self.edit_widgets['scheduled_station_ae_title'] = ttk.Entry(self.edit_frame, width=30)
        self.edit_widgets['scheduled_station_ae_title'].grid(row=6, column=1, sticky=tk.EW, padx=5, pady=5)
        self.edit_widgets['scheduled_station_ae_title'].bind("<KeyRelease>", self.on_field_change)
        
        # Configure column weights for proper resizing
        for i in range(4):
            self.edit_frame.columnconfigure(i, weight=1)

    def on_field_change(self, event=None):
        """FIXED: Detect changes and enable save button"""
        if not self.current_editing_id:
            return
        
        # Check if any field has changed from original
        current_data = self.get_current_edit_data()
        original_data = self.original_data.get(self.current_editing_id, {})
        
        changes_detected = False
        for field, current_value in current_data.items():
            if field != 'modality':  # Skip modality as it's read-only
                original_value = str(original_data.get(field, ""))
                if str(current_value) != original_value:
                    changes_detected = True
                    break
        
        if changes_detected != self.changes_made:
            self.changes_made = changes_detected
            self.save_button.config(state=tk.NORMAL if changes_detected else tk.DISABLED)

    def get_current_edit_data(self):
        """Get current data from edit widgets"""
        data = {}
        for field, widget in self.edit_widgets.items():
            if isinstance(widget, ttk.Combobox):
                data[field] = widget.get()
            else:
                data[field] = widget.get()
        return data

    def on_item_double_click(self, event):
        """FIXED: Handle double-click to edit record"""
        selection = self.tree.selection()
        if not selection:
            return
        
        item = self.tree.item(selection[0])
        values = item['values']
        
        if len(values) < 14:
            return
        
        # Store original data
        record_id = values[0]
        self.current_editing_id = record_id
        
        # Get full record from database
        record = get_patient_record_by_db_id(record_id)
        if not record:
            messagebox.showerror("Error", "Could not load record from database.")
            return
        
        self.original_data[record_id] = record.copy()
        
        # Populate edit widgets
        field_mapping = {
            'patient_name': record.get('patient_name', ''),
            'patient_id': record.get('patient_id', ''),
            'accession_number': record.get('accession_number', ''),
            'dob_yyyymmdd': record.get('dob_yyyymmdd', ''),
            'sex': record.get('sex', ''),
            'study_date': record.get('study_date', ''),
            'study_time': record.get('study_time', ''),
            'study_description': record.get('study_description', ''),
            'referred_from': record.get('referred_from', ''),
            'modality': record.get('modality', ''),
            'requesting_physician': record.get('requesting_physician', ''),
            'requested_procedure_id': record.get('requested_procedure_id', ''),
            'scheduled_station_ae_title': record.get('scheduled_station_ae_title', '')
        }
        
        for field, value in field_mapping.items():
            widget = self.edit_widgets[field]
            if isinstance(widget, ttk.Combobox):
                widget.set(str(value))
            else:
                if widget['state'] != 'readonly':
                    widget.delete(0, tk.END)
                    widget.insert(0, str(value))
                else:
                    # For readonly widgets, temporarily enable them to set value
                    widget.config(state='normal')
                    widget.delete(0, tk.END)
                    widget.insert(0, str(value))
                    widget.config(state='readonly')
        
        # Show edit frame
        self.edit_frame.pack(fill=tk.X, pady=(10,0))
        self.changes_made = False
        self.save_button.config(state=tk.DISABLED)

    def save_changes(self):
        """FIXED: Save changes to database"""
        if not self.current_editing_id or not self.changes_made:
            return
        
        current_data = self.get_current_edit_data()
        
        # Validate required fields
        required_fields = ['patient_name', 'patient_id', 'accession_number', 'dob_yyyymmdd', 'sex', 'study_description']
        for field in required_fields:
            if not current_data.get(field, '').strip():
                messagebox.showerror("Validation Error", f"{field.replace('_', ' ').title()} is required.")
                return
        
        # Prepare data for database update
        update_data = {
            "Patient Name": current_data['patient_name'],
            "Patient ID": current_data['patient_id'],
            "Accession Number": current_data['accession_number'],
            "Date of Birth": current_data['dob_yyyymmdd'],
            "Sex": current_data['sex'],
            "Study Date": current_data['study_date'],
            "Study Time": current_data['study_time'],
            "Study Description": current_data['study_description'],
            "Referred From": current_data['referred_from'],
            "Modality": current_data['modality'],
            "Requesting Physician": current_data['requesting_physician'],
            "Requested Procedure ID": current_data['requested_procedure_id'],
            "Scheduled Station AE Title": current_data['scheduled_station_ae_title']
        }
        
        try:
            success = update_patient_record_db(self.current_editing_id, update_data)
            if success:
                messagebox.showinfo("Success", "Record updated successfully.")
                self.changes_made = False
                self.save_button.config(state=tk.DISABLED)
                self.refresh_list()
                self.edit_frame.pack_forget()
                self.current_editing_id = None
                logging.info(f"Successfully updated patient record ID: {self.current_editing_id}")
            else:
                messagebox.showerror("Error", "Failed to update record in database.")
                logging.error(f"Failed to update patient record ID: {self.current_editing_id}")
        except Exception as e:
            messagebox.showerror("Database Error", f"Error updating record: {e}")
            logging.exception(f"Error updating patient record ID: {self.current_editing_id}")

    def refresh_list(self, event=None):
        """Refresh the worklist display"""
        self.tree.delete(*self.tree.get_children())
        
        try:
            all_data, _ = get_all_patient_records_db("")
            for record in all_data:
                values = [
                    record.get('id', ''),
                    record.get('patient_name', ''),
                    record.get('patient_id', ''),
                    record.get('accession_number', ''),
                    record.get('dob_yyyymmdd', ''),
                    record.get('sex', ''),
                    record.get('study_date', ''),
                    record.get('study_time', ''),
                    record.get('study_description', ''),
                    record.get('referred_from', ''),
                    record.get('modality', ''),
                    record.get('requesting_physician', ''),
                    record.get('requested_procedure_id', ''),
                    record.get('scheduled_station_ae_title', '')
                ]
                item_id = self.tree.insert("", tk.END, values=values)
                # Tag modality column as readonly
                self.tree.set(item_id, "modality", values[10])
                
            logging.info(f"Loaded {len(all_data)} records into worklist viewer.")
            
        except Exception as e:
            messagebox.showerror("Database Error", f"Error loading worklist data: {e}", parent=self)
            logging.exception("Error loading worklist data.")
        
        # Hide edit frame if visible
        self.edit_frame.pack_forget()
        self.current_editing_id = None
        self.changes_made = False
        self.save_button.config(state=tk.DISABLED)

    def on_closing(self):

           """Handle window closing"""
           if self.changes_made:
            if messagebox.askyesno("Unsaved Changes", "You have unsaved changes. Do you want to save them before closing?"):
                self.save_changes()
                self.destroy()

    def delete_selected_records(self):
        """Delete the selected patient records (supports multiple selection)"""
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select one or more records to delete.", parent=self)
            return
        
        # Collect record information
        records_to_delete = []
        for item_id in selection:
            item = self.tree.item(item_id)
            values = item['values']
            
            if len(values) >= 3:
                record_id = values[0]
                patient_name = values[1]
                patient_id = values[2]
                records_to_delete.append({
                    'id': record_id,
                    'name': patient_name,
                    'patient_id': patient_id
                })
        
        if not records_to_delete:
            return
        
        # Confirm deletion
        if len(records_to_delete) == 1:
            record = records_to_delete[0]
            confirm_msg = (f"Are you sure you want to delete this patient record?\n\n"
                          f"Patient: {record['name']} ({record['patient_id']})\n"
                          f"This action cannot be undone.")
        else:
            patient_list = "\n".join([f"• {r['name']} ({r['patient_id']})" for r in records_to_delete[:5]])
            if len(records_to_delete) > 5:
                patient_list += f"\n... and {len(records_to_delete) - 5} more"
            
            confirm_msg = (f"Are you sure you want to delete {len(records_to_delete)} patient records?\n\n"
                          f"Records to delete:\n{patient_list}\n\n"
                          f"This action cannot be undone.")
        
        if messagebox.askyesno("Confirm Delete", confirm_msg, parent=self):
            deleted_count = 0
            failed_count = 0
            
            for record in records_to_delete:
                try:
                    # Delete from database
                    query = "DELETE FROM patient_records WHERE id = ?"
                    success = db_execute(query, (record['id'],), commit=True)
                    
                    if success:
                        deleted_count += 1
                        logging.info(f"Successfully deleted patient record ID: {record['id']} ({record['name']})")
                    else:
                        failed_count += 1
                        logging.error(f"Failed to delete patient record ID: {record['id']} ({record['name']})")
                except Exception as e:
                    failed_count += 1
                    logging.exception(f"Error deleting patient record ID: {record['id']} ({record['name']}): {e}")
            
            # Show results
            if failed_count == 0:
                if deleted_count == 1:
                    messagebox.showinfo("Success", "Patient record deleted successfully.", parent=self)
                else:
                    messagebox.showinfo("Success", f"{deleted_count} patient records deleted successfully.", parent=self)
            else:
                if deleted_count > 0:
                    messagebox.showwarning("Partial Success", 
                                         f"{deleted_count} records deleted successfully.\n"
                                         f"{failed_count} records failed to delete.\n"
                                         f"Check logs for details.", parent=self)
                else:
                    messagebox.showerror("Error", 
                                       f"Failed to delete {failed_count} records.\n"
                                       f"Check logs for details.", parent=self)
            
            # Refresh the list
            self.refresh_list()

# --- Main Application Entry Point ---
def main():
    root = tk.Tk()
    app = PatientRegistrationApp(root)
    
    def on_closing():
        logging.info("Application closing initiated by user.")
        app.shutdown()
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    try:
        root.mainloop()
    except KeyboardInterrupt:
        logging.info("Application interrupted by keyboard.")
        app.shutdown()
    except Exception as e:
        logging.exception("Unexpected error in main loop.")
        app.shutdown()
        raise
    finally:
        logging.info("Application main loop ended.")

if __name__ == "__main__":
    main()
